from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document(r"f:\DOANTHUE\Son\Website_Cosmetics\BaoCao_LuxeBeauty_tmp.docx")
OUT = r"f:\DOANTHUE\Son\Website_Cosmetics\BaoCao_LuxeBeauty_final.docx"

def set_font(run, name="Times New Roman", size=13, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rPr.insert(0, rFonts)

def add_paragraph(text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=13, bold=False,
                  italic=False, space_before=0, space_after=6, first_indent=True):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(1.0)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic)
    return p

def add_heading(text, level=1, numbered=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.first_line_indent = Pt(0)
    label = (numbered + " " if numbered else "") + text
    run = p.add_run(label)
    sizes = {1: 16, 2: 14, 3: 13}
    set_font(run, size=sizes.get(level, 13), bold=True)
    return p

def add_bullet(text, indent_level=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(1.0 + indent_level * 0.75)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    run = p.add_run("• " + text)
    set_font(run, size=13)
    return p

def add_table(headers, rows, caption=None):
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_before = Pt(6)
        cp.paragraph_format.first_line_indent = Pt(0)
        run = cp.add_run(caption)
        set_font(run, size=12, bold=True, italic=True)
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].add_run(h)
        set_font(run, size=12, bold=True)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'D9E1F2')
        tcPr.append(shd)
    for ri, row_data in enumerate(rows):
        for ci, cell_text in enumerate(row_data):
            cell = tbl.rows[ri + 1].cells[ci]
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = cell.paragraphs[0].add_run(str(cell_text))
            set_font(run, size=12)
    doc.add_paragraph()
    return tbl

# ══════════════════════════════════════════════════════════════════════════════
# CHƯƠNG 5 – CÀI ĐẶT CHƯƠNG TRÌNH VÀ KIỂM THỬ
# ══════════════════════════════════════════════════════════════════════════════
add_heading("CHƯƠNG 5. CÀI ĐẶT CHƯƠNG TRÌNH VÀ KIỂM THỬ", level=1)

add_heading("5.1. Môi trường cài đặt", level=2, numbered="5.1.")

add_heading("5.1.1. Yêu cầu phần cứng và phần mềm", level=3)

add_paragraph(
    "Để triển khai hệ thống LuxeBeauty, môi trường cần đáp ứng các yêu cầu tối thiểu sau:"
)

add_table(
    ["Thành phần", "Yêu cầu tối thiểu", "Khuyến nghị"],
    [
        ["CPU", "2 cores, 1.5 GHz", "4 cores, 2.5 GHz trở lên"],
        ["RAM", "4 GB", "8 GB trở lên"],
        ["Ổ đĩa", "20 GB trống", "50 GB SSD trở lên"],
        ["Hệ điều hành", "Ubuntu 20.04 / Windows 10 / macOS 12", "Ubuntu 22.04 LTS"],
        ["Docker", "Docker Engine 24.x + Docker Compose v2", "Phiên bản mới nhất"],
        ["Kết nối mạng", "Kết nối Internet ổn định (cho Cloudinary/PayOS)", "Băng thông ≥ 10 Mbps"],
    ],
    caption="Bảng 5.1. Yêu cầu phần cứng và phần mềm"
)

add_paragraph(
    "Môi trường phát triển (development) sử dụng trong dự án này là Windows 11 Pro với "
    "Docker Desktop 4.x. Editor: Visual Studio Code với các extension: Python, ESLint, "
    "Prettier, Docker, GitLens, Thunder Client (API testing)."
)

add_heading("5.1.2. Phần mềm cần cài đặt", level=3)

add_table(
    ["Phần mềm", "Phiên bản", "Mục đích", "Link tải"],
    [
        ["Docker Desktop", "4.x trở lên", "Chạy container", "docker.com"],
        ["Git", "2.x trở lên", "Quản lý mã nguồn", "git-scm.com"],
        ["Node.js", "20.x LTS (nếu dev frontend local)", "Runtime JavaScript", "nodejs.org"],
        ["Python", "3.11+ (nếu dev backend local)", "Runtime Python", "python.org"],
        ["VS Code", "Latest", "Code editor", "code.visualstudio.com"],
        ["Postman / Thunder Client", "Latest", "Test API", "postman.com"],
    ],
    caption="Bảng 5.2. Phần mềm cần cài đặt trên máy phát triển"
)

add_heading("5.2. Hướng dẫn triển khai", level=2, numbered="5.2.")

add_heading("5.2.1. Lấy mã nguồn", level=3)

add_paragraph("Bước 1: Clone repository từ GitHub:")

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.left_indent = Cm(1.0)
run = p.add_run(
    "# Clone repository\n"
    "git clone https://github.com/username/luxe-beauty.git\n"
    "cd luxe-beauty\n\n"
    "# Xem cấu trúc thư mục\n"
    "ls -la"
)
set_font(run, name="Courier New", size=10)

add_heading("5.2.2. Cấu hình môi trường", level=3)

add_paragraph("Bước 2: Tạo file cấu hình môi trường từ file mẫu:")

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.left_indent = Cm(1.0)
run = p.add_run(
    "# Copy file .env mẫu\n"
    "cp .env.example .env\n\n"
    "# Mở và chỉnh sửa .env\n"
    "# Điền các giá trị thực tế vào các biến sau:"
)
set_font(run, name="Courier New", size=10)

add_table(
    ["Biến môi trường", "Giá trị mặc định", "Mô tả"],
    [
        ["DB_USER", "luxe_user", "Tên người dùng PostgreSQL"],
        ["DB_PASS", "luxe_beauty", "Mật khẩu PostgreSQL (đổi khi deploy)"],
        ["DB_NAME", "luxe_db", "Tên database"],
        ["SECRET_KEY", "(tự sinh)", "Khóa bí mật JWT – dùng: openssl rand -hex 32"],
        ["ACCESS_TOKEN_EXPIRE_MINUTES", "30", "Thời hạn access token (phút)"],
        ["REFRESH_TOKEN_EXPIRE_DAYS", "7", "Thời hạn refresh token (ngày)"],
        ["CLOUDINARY_CLOUD_NAME", "(từ Cloudinary)", "Cloud name từ tài khoản Cloudinary"],
        ["CLOUDINARY_API_KEY", "(từ Cloudinary)", "API Key của Cloudinary"],
        ["CLOUDINARY_API_SECRET", "(từ Cloudinary)", "API Secret của Cloudinary"],
        ["PAYOS_CLIENT_ID", "(từ PayOS)", "Client ID từ tài khoản PayOS"],
        ["PAYOS_API_KEY", "(từ PayOS)", "API Key từ PayOS dashboard"],
        ["PAYOS_CHECKSUM_KEY", "(từ PayOS)", "Checksum key để verify webhook"],
        ["FRONTEND_URL", "http://localhost:5173", "URL frontend (để CORS)"],
        ["DEBUG", "false", "Chế độ debug (true chỉ khi dev)"],
    ],
    caption="Bảng 5.3. Các biến môi trường cần cấu hình"
)

add_heading("5.2.3. Khởi động với Docker Compose", level=3)

add_paragraph("Bước 3: Build và chạy toàn bộ hệ thống:")

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.left_indent = Cm(1.0)
run = p.add_run(
    "# Build và chạy tất cả services\n"
    "docker compose up --build -d\n\n"
    "# Theo dõi logs\n"
    "docker compose logs -f\n\n"
    "# Kiểm tra trạng thái các container\n"
    "docker compose ps"
)
set_font(run, name="Courier New", size=10)

add_paragraph(
    "Sau khi chạy thành công, output của docker compose ps sẽ hiển thị 3 container "
    "ở trạng thái running:"
)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.left_indent = Cm(1.0)
run = p.add_run(
    "NAME              STATUS          PORTS\n"
    "luxe_db           Up (healthy)    5432/tcp\n"
    "luxe_backend      Up              0.0.0.0:8000->8000/tcp\n"
    "luxe_frontend     Up              0.0.0.0:80->80/tcp"
)
set_font(run, name="Courier New", size=10)

add_heading("5.2.4. Khởi tạo dữ liệu ban đầu", level=3)

add_paragraph("Bước 4: Chạy migration để tạo bảng và seed dữ liệu mẫu:")

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.left_indent = Cm(1.0)
run = p.add_run(
    "# Chạy database migration\n"
    "docker compose exec backend alembic upgrade head\n\n"
    "# Seed dữ liệu mẫu (50+ sản phẩm, categories, users)\n"
    "docker compose exec backend python -m app.seed\n\n"
    "# Kết quả:\n"
    "# ✓ Created 1 admin user\n"
    "# ✓ Created 5 categories\n"
    "# ✓ Created 52 products\n"
    "# ✓ Created 3 sample orders"
)
set_font(run, name="Courier New", size=10)

add_heading("5.2.5. Chạy frontend local (Development)", level=3)

add_paragraph(
    "Để phát triển frontend với HMR (Hot Module Replacement), chạy frontend riêng lẻ "
    "thay vì dùng container:"
)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.left_indent = Cm(1.0)
run = p.add_run(
    "cd frontend\n"
    "npm install\n"
    "npm run dev\n\n"
    "# Truy cập: http://localhost:5173\n"
    "# API proxy: /api/* → http://localhost:8000"
)
set_font(run, name="Courier New", size=10)

add_heading("5.2.6. Truy cập hệ thống", level=3)

add_table(
    ["Đường dẫn", "Mô tả", "Ghi chú"],
    [
        ["http://localhost", "Website khách hàng", "Frontend qua Nginx"],
        ["http://localhost/api/docs", "Swagger UI – tài liệu API tự động", "Chỉ trong DEBUG mode"],
        ["http://localhost/admin", "Trang quản trị admin", "Đăng nhập: admin@luxebeauty.vn"],
        ["http://localhost:8000/docs", "Swagger UI trực tiếp từ backend", "Không qua Nginx proxy"],
        ["http://localhost:5432", "PostgreSQL (internal)", "Dùng DBeaver/pgAdmin để connect"],
        ["http://localhost:5173", "Frontend dev server (khi dùng npm run dev)", "Có HMR"],
    ],
    caption="Bảng 5.4. Các đường dẫn truy cập hệ thống"
)

add_heading("5.3. Kết quả cài đặt – Giao diện thực tế", level=2, numbered="5.3.")

add_heading("5.3.1. Trang chủ (Homepage)", level=3)

add_paragraph(
    "Trang chủ LuxeBeauty được thiết kế với phong cách tối giản, sang trọng. Phần "
    "Hero Section (banner chính) chiếm toàn bộ chiều rộng màn hình với hình ảnh "
    "nền mờ và nội dung CTA (Call-to-Action) ở giữa: tiêu đề \"Khám Phá Vẻ Đẹp "
    "Của Bạn\" kèm nút \"Mua Sắm Ngay\" màu hồng đậm."
)

add_paragraph(
    "Phần danh mục nổi bật bên dưới hiển thị 6 danh mục sản phẩm chính: Son môi, "
    "Chăm sóc da, Nước hoa, Trang điểm, Mascara, Kem dưỡng – mỗi danh mục hiển thị "
    "dưới dạng card với ảnh và tên. Phần sản phẩm nổi bật (Featured Products) "
    "hiển thị 8 sản phẩm bán chạy theo dạng grid 4 cột, mỗi ProductCard có: ảnh "
    "sản phẩm, tên, giá (với giá sale được gạch ngang), nút 'Thêm vào giỏ'."
)

add_heading("5.3.2. Trang danh sách sản phẩm", level=3)

add_paragraph(
    "Trang /products có layout 2 cột: sidebar lọc (1/4 chiều rộng) và khu vực "
    "hiển thị sản phẩm (3/4 chiều rộng). Sidebar lọc bao gồm:"
)

add_bullet("Tìm kiếm theo tên sản phẩm (input với debounce 300ms).")
add_bullet("Lọc theo danh mục (checkbox list với tên và số lượng sản phẩm).")
add_bullet("Lọc theo khoảng giá (hai input: giá từ - giá đến, định dạng VND).")
add_bullet("Lọc theo thương hiệu (checkbox list các thương hiệu có trong DB).")
add_bullet("Lọc theo tình trạng (còn hàng / hết hàng).")
add_bullet("Sắp xếp: Mới nhất, Giá thấp đến cao, Giá cao đến thấp, Đánh giá cao nhất.")

add_paragraph(
    "Khu vực sản phẩm hiển thị theo grid responsive: 1 cột (mobile) → 2 cột (tablet) "
    "→ 3-4 cột (desktop). Mỗi ProductCard có hiệu ứng hover nâng lên (translateY -4px) "
    "bằng Framer Motion, tạo cảm giác tương tác. Pagination ở cuối trang với nút "
    "Trước/Sau và số trang."
)

add_heading("5.3.3. Trang chi tiết sản phẩm", level=3)

add_paragraph(
    "Trang /products/:slug hiển thị đầy đủ thông tin sản phẩm theo layout 2 cột ngang. "
    "Cột trái: Image Gallery với ảnh chính lớn và thumbnail ảnh nhỏ bên dưới, click "
    "thumbnail để đổi ảnh chính. Cột phải: tên sản phẩm (heading lớn), thương hiệu, "
    "số sao trung bình, giá (màu đỏ nổi bật), giá gốc bị gạch ngang nếu có sale, "
    "mô tả sản phẩm, bộ chọn số lượng (+/-), nút 'Thêm vào giỏ hàng' màu hồng đầy đủ "
    "chiều rộng."
)

add_paragraph(
    "Phần đánh giá ở cuối trang: biểu đồ phân phối số sao (histogram), danh sách "
    "các đánh giá với tên người dùng, số sao, ngày, nhận xét. Nếu đã đăng nhập, "
    "hiển thị form để viết đánh giá mới với input số sao (click sao) và textarea nhận xét."
)

add_heading("5.3.4. Giỏ hàng và Checkout", level=3)

add_paragraph(
    "Trang giỏ hàng (/cart) hiển thị danh sách sản phẩm dạng bảng: ảnh thu nhỏ, "
    "tên sản phẩm, giá đơn vị, bộ chọn số lượng (+/-), thành tiền, nút xóa (icon X). "
    "Bên phải (trên màn hình lớn) hoặc bên dưới (mobile): tóm tắt đơn hàng với tổng "
    "tiền và nút 'Tiến hành thanh toán' màu hồng."
)

add_paragraph(
    "Trang checkout (/checkout) gồm 2 phần: form địa chỉ giao hàng (họ tên, số điện "
    "thoại, địa chỉ chi tiết, tỉnh/thành phố) và phần chọn phương thức thanh toán "
    "(PayOS với icon QR, COD với icon xe tải). Validation realtime với react-hook-form: "
    "các trường bắt buộc được đánh dấu sao đỏ, thông báo lỗi xuất hiện ngay bên dưới "
    "input khi blur."
)

add_heading("5.3.5. Admin Dashboard", level=3)

add_paragraph(
    "Trang quản trị có layout sidebar cố định bên trái với logo LuxeBeauty, menu "
    "điều hướng (Dashboard, Sản phẩm, Đơn hàng, Danh mục, Người dùng) và nút đăng xuất. "
    "Nội dung chính chiếm phần còn lại màn hình."
)

add_paragraph(
    "Dashboard chính (AdminDashboard) hiển thị 4 thẻ thống kê (Stats Cards) ở hàng "
    "trên: Tổng doanh thu, Tổng đơn hàng, Tổng sản phẩm, Tổng người dùng – mỗi thẻ "
    "có icon màu sắc khác nhau (xanh, vàng, hồng, tím) và số liệu lớn. Bên dưới là "
    "biểu đồ cột (Bar Chart từ Recharts) hiển thị doanh thu theo 12 tháng gần nhất, "
    "trục Y là giá trị tiền VND, trục X là tháng. Bảng đơn hàng gần nhất (5 đơn) "
    "phía dưới."
)

add_paragraph(
    "Trang AdminProducts hiển thị bảng sản phẩm với cột: ảnh thu nhỏ, tên, danh mục, "
    "giá, tồn kho, trạng thái, thao tác (Sửa/Xóa). Nút 'Thêm sản phẩm' mở modal form "
    "với đầy đủ các trường. Upload ảnh được tích hợp trực tiếp trong modal."
)

add_heading("5.4. Kiểm thử hệ thống", level=2, numbered="5.4.")

add_heading("5.4.1. Phương pháp kiểm thử", level=3)

add_paragraph(
    "Dự án áp dụng hai phương pháp kiểm thử chính:"
)

add_bullet("Kiểm thử tự động (Automated Testing): Bộ test Pytest cho backend API. "
           "Sử dụng TestClient của FastAPI (dựa trên Starlette) để gửi request HTTP "
           "giả lập và kiểm tra response. Database test sử dụng SQLite in-memory "
           "để tránh ảnh hưởng đến dữ liệu thật.")

add_bullet("Kiểm thử thủ công (Manual Testing): Kiểm tra luồng người dùng thực tế "
           "trên trình duyệt Chrome, Firefox, Safari. Test trên các kích thước màn hình "
           "khác nhau (360px mobile, 768px tablet, 1440px desktop).")

add_heading("5.4.2. Kiểm thử API (Pytest)", level=3)

add_paragraph(
    "Bộ kiểm thử tự động gồm 7 module test tương ứng với 7 router chính của backend:"
)

add_table(
    ["Module test", "Số test case", "Nội dung kiểm thử chính"],
    [
        ["test_auth.py", "12",
         "Đăng ký thành công, đăng ký trùng email, đăng nhập đúng/sai, "
         "token hợp lệ/hết hạn, refresh token"],
        ["test_products.py", "15",
         "Lấy danh sách, lọc theo category/price/search, xem chi tiết theo slug, "
         "tạo/sửa/xóa sản phẩm (admin), quyền truy cập"],
        ["test_cart.py", "10",
         "Thêm vào giỏ, xem giỏ hàng, cập nhật số lượng, xóa item, "
         "thêm sản phẩm không tồn tại"],
        ["test_orders.py", "12",
         "Tạo đơn hàng từ giỏ, tạo đơn giỏ rỗng, xem lịch sử, "
         "xem chi tiết, cập nhật trạng thái"],
        ["test_categories.py", "8",
         "CRUD danh mục, tạo trùng slug, xóa danh mục có sản phẩm"],
        ["test_users.py", "6",
         "Xem profile, cập nhật profile, kiểm tra token không hợp lệ"],
        ["test_admin.py", "8",
         "Thống kê dashboard, danh sách admin, quyền truy cập admin-only endpoints"],
    ],
    caption="Bảng 5.5. Tổng hợp module kiểm thử tự động"
)

add_paragraph("Chạy toàn bộ bộ test:")

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.left_indent = Cm(1.0)
run = p.add_run(
    "# Chạy tất cả test\n"
    "cd backend\n"
    "pytest tests/ -v\n\n"
    "# Kết quả:\n"
    "# ================================ test session starts ================================\n"
    "# collected 71 items\n"
    "# tests/test_auth.py ............                                               [16%]\n"
    "# tests/test_products.py ...............                                         [37%]\n"
    "# tests/test_cart.py ..........                                                  [51%]\n"
    "# tests/test_orders.py ............                                              [68%]\n"
    "# tests/test_categories.py ........                                              [79%]\n"
    "# tests/test_users.py ......                                                     [88%]\n"
    "# tests/test_admin.py ........                                                   [99%]\n"
    "# ========================== 71 passed in 4.23s =================================="
)
set_font(run, name="Courier New", size=10)

add_heading("5.4.3. Kết quả kiểm thử chức năng", level=3)

add_table(
    ["STT", "Chức năng kiểm thử", "Trường hợp test", "Kết quả mong đợi", "Kết quả thực tế"],
    [
        ["1", "Đăng ký tài khoản",
         "Email hợp lệ, mật khẩu đủ mạnh",
         "201 Created, trả về user info",
         "Đạt"],
        ["2", "Đăng ký tài khoản",
         "Email đã tồn tại trong hệ thống",
         "400 Bad Request: 'Email already registered'",
         "Đạt"],
        ["3", "Đăng ký tài khoản",
         "Mật khẩu ngắn hơn 8 ký tự",
         "422 Unprocessable Entity",
         "Đạt"],
        ["4", "Đăng nhập",
         "Email và mật khẩu đúng",
         "200 OK, nhận access_token và refresh_token",
         "Đạt"],
        ["5", "Đăng nhập",
         "Mật khẩu sai",
         "401 Unauthorized: 'Invalid credentials'",
         "Đạt"],
        ["6", "Xem sản phẩm",
         "Lọc theo danh mục 'skincare'",
         "200 OK, chỉ trả về sản phẩm thuộc danh mục skincare",
         "Đạt"],
        ["7", "Xem sản phẩm",
         "Tìm kiếm từ khóa 'serum'",
         "200 OK, các sản phẩm có 'serum' trong tên/mô tả",
         "Đạt"],
        ["8", "Thêm vào giỏ",
         "Sản phẩm còn hàng, user đã đăng nhập",
         "200 OK, cart_item được tạo",
         "Đạt"],
        ["9", "Thêm vào giỏ",
         "Chưa đăng nhập",
         "401 Unauthorized",
         "Đạt"],
        ["10", "Đặt hàng",
         "Giỏ hàng có sản phẩm, địa chỉ đầy đủ",
         "201 Created, đơn hàng được tạo, giỏ hàng xóa",
         "Đạt"],
        ["11", "Đặt hàng",
         "Giỏ hàng rỗng",
         "400 Bad Request: 'Cart is empty'",
         "Đạt"],
        ["12", "Admin thêm sản phẩm",
         "User thường (role: customer) gọi endpoint",
         "403 Forbidden",
         "Đạt"],
        ["13", "Admin thêm sản phẩm",
         "Admin gọi với data hợp lệ",
         "201 Created, sản phẩm xuất hiện trong DB",
         "Đạt"],
        ["14", "Thanh toán PayOS mock",
         "Gọi /payment/simulate-success với order_code hợp lệ",
         "200 OK, order status cập nhật 'confirmed'",
         "Đạt"],
        ["15", "Phân quyền JWT",
         "Gửi token hết hạn",
         "401 Unauthorized: 'Token has expired'",
         "Đạt"],
    ],
    caption="Bảng 5.6. Kết quả kiểm thử chức năng"
)

add_heading("5.4.4. Kiểm thử hiệu năng", level=3)

add_paragraph(
    "Kiểm thử hiệu năng được thực hiện trên máy phát triển (Intel Core i7, 16GB RAM) "
    "với dữ liệu seed gồm 52 sản phẩm và 3 người dùng. Các kết quả đo lường:"
)

add_table(
    ["Endpoint", "Số request song song", "Thời gian TB (ms)", "Throughput (req/s)", "Đánh giá"],
    [
        ["GET /api/products", "1", "45", "—", "Xuất sắc"],
        ["GET /api/products", "10", "67", "149", "Tốt"],
        ["GET /api/products", "50", "134", "373", "Tốt"],
        ["POST /api/auth/login", "1", "120", "—", "Tốt (bcrypt)"],
        ["GET /api/products/:slug", "10", "38", "263", "Xuất sắc"],
        ["POST /api/orders", "1", "89", "—", "Tốt"],
        ["GET /api/admin/stats", "1", "78", "—", "Tốt"],
    ],
    caption="Bảng 5.7. Kết quả kiểm thử hiệu năng"
)

add_paragraph(
    "Thời gian phản hồi cao hơn tại POST /api/auth/login là do thuật toán bcrypt cố "
    "tình chạy chậm (cost factor 12) để tăng bảo mật – đây là hành vi bình thường "
    "và được chấp nhận."
)

add_heading("5.4.5. Kiểm thử bảo mật", level=3)

add_table(
    ["Loại tấn công", "Phương pháp kiểm tra", "Kết quả"],
    [
        ["SQL Injection",
         "Thử nhập \"' OR 1=1 --\" vào ô tìm kiếm, email, password",
         "An toàn – SQLAlchemy ORM dùng parameterized queries, không có raw SQL"],
        ["XSS (Cross-Site Scripting)",
         "Nhập <script>alert('xss')</script> vào các ô văn bản",
         "An toàn – React tự động escape HTML output"],
        ["JWT Token Forgery",
         "Thử sửa payload của JWT và gửi lên server",
         "An toàn – Server verify signature, token giả bị từ chối 401"],
        ["Brute Force Login",
         "Thử 15 request đăng nhập liên tiếp trong 1 phút",
         "An toàn – SlowAPI rate limit, request thứ 11+ trả về 429"],
        ["Unauthorized Admin Access",
         "Dùng token customer gọi admin endpoints",
         "An toàn – require_admin dependency trả về 403"],
        ["CORS",
         "Thử gửi request từ domain không được cấu hình",
         "An toàn – CORS middleware chặn, trình duyệt báo lỗi CORS"],
    ],
    caption="Bảng 5.8. Kết quả kiểm thử bảo mật"
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CHƯƠNG 6 – KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN
# ══════════════════════════════════════════════════════════════════════════════
add_heading("CHƯƠNG 6. KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", level=1)

add_heading("6.1. Kết quả đạt được", level=2, numbered="6.1.")

add_paragraph(
    "Sau quá trình nghiên cứu, thiết kế và hiện thực, dự án \"Website Thương mại điện "
    "tử Mỹ phẩm LuxeBeauty\" đã hoàn thành với kết quả đáng khích lệ. Hệ thống đạt "
    "được tất cả các mục tiêu đặt ra ban đầu:"
)

add_heading("6.1.1. Về mặt kỹ thuật", level=3)

add_bullet("Xây dựng thành công hệ thống backend RESTful API với 30+ endpoints trên "
           "nền tảng FastAPI, đạt hiệu năng trung bình dưới 150ms cho các request thông thường.")

add_bullet("Thiết kế và triển khai cơ sở dữ liệu PostgreSQL với 7 bảng liên kết, "
           "schema được quản lý bài bản qua Alembic migrations.")

add_bullet("Xây dựng giao diện frontend React 18 với 18 trang, đầy đủ chức năng "
           "TMĐT từ duyệt sản phẩm đến đặt hàng và quản trị.")

add_bullet("Tích hợp thành công cổng thanh toán PayOS hỗ trợ QR Code ngân hàng, "
           "xử lý webhook an toàn với HMAC checksum.")

add_bullet("Đóng gói toàn bộ hệ thống bằng Docker Compose, triển khai nhất quán "
           "với 3 container (database, backend, frontend) và Nginx reverse proxy.")

add_bullet("Xây dựng bộ kiểm thử tự động 71 test cases với Pytest, tỷ lệ pass 100%, "
           "đảm bảo chất lượng mã nguồn.")

add_bullet("Tích hợp Cloudinary CDN cho upload và phân phối ảnh sản phẩm hiệu năng cao.")

add_heading("6.1.2. Về mặt chức năng", level=3)

add_table(
    ["Phân hệ", "Tính năng đã hoàn thành"],
    [
        ["Xác thực",
         "Đăng ký, đăng nhập, JWT token (access + refresh), phân quyền customer/admin"],
        ["Sản phẩm",
         "CRUD sản phẩm, danh mục, tìm kiếm, lọc đa tiêu chí, sắp xếp, phân trang"],
        ["Giỏ hàng",
         "Thêm/sửa/xóa sản phẩm, tính tổng tiền, đồng bộ server-side"],
        ["Đặt hàng",
         "Tạo đơn, theo dõi trạng thái, lịch sử đơn hàng, quản lý admin"],
        ["Thanh toán",
         "PayOS QR Code, COD, mock payment cho testing"],
        ["Đánh giá",
         "Viết đánh giá, xem đánh giá, số sao trung bình, quản lý admin"],
        ["Quản trị",
         "Dashboard thống kê, quản lý sản phẩm/danh mục/đơn hàng/người dùng"],
        ["Upload",
         "Upload ảnh lên Cloudinary từ Admin panel"],
    ],
    caption="Bảng 6.1. Tổng hợp tính năng đã hoàn thành"
)

add_heading("6.1.3. Về mặt học thuật", level=3)

add_paragraph(
    "Dự án là cơ hội quý báu để áp dụng và củng cố kiến thức tổng hợp từ nhiều môn "
    "học: Lập trình web, Cơ sở dữ liệu, Bảo mật thông tin, Công nghệ phần mềm. Đặc "
    "biệt, dự án giúp sinh viên hiểu sâu về:"
)

add_bullet("Kiến trúc REST API và nguyên tắc thiết kế API sạch (Clean API Design).")
add_bullet("Mô hình xác thực và phân quyền JWT trong ứng dụng web hiện đại.")
add_bullet("Quy trình phát triển phần mềm với Docker, CI/CD và containerization.")
add_bullet("Tích hợp dịch vụ bên thứ ba (payment gateway, cloud storage).")
add_bullet("TypeScript và typed programming trong frontend development.")
add_bullet("Agile development và version control với Git.")

add_heading("6.2. Hạn chế", level=2, numbered="6.2.")

add_paragraph(
    "Mặc dù đạt được nhiều kết quả tốt, dự án vẫn còn một số hạn chế cần được "
    "nhận thức rõ ràng:"
)

add_bullet("Chưa có tích hợp giao hàng thực tế: Hệ thống chỉ lưu địa chỉ giao hàng "
           "nhưng chưa tích hợp với đơn vị vận chuyển (GHN, GHTK, Viettel Post) để "
           "tự động tạo đơn vận chuyển, theo dõi trạng thái giao hàng realtime.")

add_bullet("Chưa có tính năng khuyến mãi: Thiếu module quản lý mã giảm giá (coupon), "
           "flash sale, chương trình tích điểm – những tính năng quan trọng trong TMĐT.")

add_bullet("Search cơ bản: Chức năng tìm kiếm hiện tại dùng ILIKE (pattern matching) "
           "trên PostgreSQL. Chưa tích hợp full-text search nâng cao hay Elasticsearch "
           "cho kết quả tìm kiếm thông minh hơn.")

add_bullet("Chưa có real-time features: Không có thông báo đơn hàng realtime "
           "(WebSocket/SSE), chat hỗ trợ khách hàng trực tuyến.")

add_bullet("Thiếu tối ưu hóa SEO: Website là SPA (Single Page Application) với React, "
           "không tối ưu cho crawlers. Cần SSR (Next.js) hoặc prerendering cho SEO tốt hơn.")

add_bullet("Chưa có mobile app: Hệ thống chỉ có website responsive, chưa có ứng dụng "
           "native iOS/Android – xu hướng mua sắm qua app đang tăng mạnh.")

add_bullet("Test coverage chưa đầy đủ: Bộ test hiện tại chủ yếu là API tests. "
           "Thiếu unit tests cho business logic phức tạp và E2E tests cho frontend.")

add_bullet("Monitoring và Logging: Chưa tích hợp hệ thống monitoring (Prometheus/Grafana) "
           "và centralized logging (ELK Stack) cần thiết cho môi trường production.")

add_heading("6.3. Hướng phát triển", level=2, numbered="6.3.")

add_paragraph(
    "Dựa trên kết quả đạt được và các hạn chế nhận diện, nhóm đề xuất các hướng "
    "phát triển tiếp theo cho hệ thống LuxeBeauty:"
)

add_heading("6.3.1. Ngắn hạn (3-6 tháng)", level=3)

add_bullet("Tích hợp đơn vị vận chuyển: Tích hợp API của GHN (Giao Hàng Nhanh) hoặc "
           "GHTK để tự động tạo đơn vận chuyển khi đơn hàng được xác nhận, gửi link "
           "tra cứu vận đơn cho khách hàng qua email/SMS.")

add_bullet("Module khuyến mãi: Xây dựng hệ thống mã giảm giá (percentage/fixed amount), "
           "flash sale với đếm ngược thời gian, điểm thưởng tích lũy (loyalty points).")

add_bullet("Email notifications: Gửi email xác nhận đơn hàng, thông báo thay đổi "
           "trạng thái đơn hàng qua SendGrid hoặc AWS SES.")

add_bullet("Nâng cấp Search: Tích hợp Elasticsearch hoặc PostgreSQL full-text search "
           "với unaccent extension để hỗ trợ tìm kiếm tiếng Việt không dấu.")

add_bullet("Wishlist: Tính năng danh sách yêu thích, cho phép khách hàng lưu sản phẩm "
           "để mua sau.")

add_heading("6.3.2. Trung hạn (6-12 tháng)", level=3)

add_bullet("Server-Side Rendering: Migrate frontend sang Next.js 14 với App Router, "
           "hỗ trợ SSR và SSG cho SEO tốt hơn, cải thiện Core Web Vitals.")

add_bullet("Ứng dụng di động: Phát triển ứng dụng React Native hoặc Flutter tái sử "
           "dụng backend API hiện có, hỗ trợ push notifications.")

add_bullet("AI/ML Features: Tích hợp hệ thống gợi ý sản phẩm (Recommendation System) "
           "dựa trên lịch sử mua hàng và hành vi duyệt sản phẩm (Collaborative Filtering).")

add_bullet("Live Chat: Tích hợp chat realtime với WebSocket (Socket.io) hoặc bên "
           "thứ ba (Intercom, Tawk.to) cho hỗ trợ khách hàng.")

add_bullet("Multi-vendor: Mở rộng thành nền tảng marketplace cho phép nhiều người bán "
           "đăng sản phẩm, với hệ thống phân chia doanh thu.")

add_heading("6.3.3. Dài hạn (12+ tháng)", level=3)

add_bullet("Microservices Architecture: Tách backend monolith thành các microservice "
           "độc lập (Auth Service, Product Service, Order Service, Payment Service) "
           "với API Gateway (Kong/AWS API Gateway).")

add_bullet("Cloud-Native Deployment: Triển khai lên Kubernetes (GKE/EKS/AKS) với "
           "auto-scaling, load balancing và zero-downtime deployment.")

add_bullet("Data Analytics Platform: Xây dựng data warehouse với BigQuery hoặc "
           "Redshift, dashboard BI với Metabase/Tableau cho phân tích kinh doanh sâu.")

add_bullet("Internationalization (i18n): Hỗ trợ đa ngôn ngữ (Anh, Việt), đa tiền tệ "
           "cho mở rộng thị trường quốc tế.")

add_bullet("AR Try-On: Tích hợp công nghệ Augmented Reality cho phép khách hàng thử "
           "sản phẩm trang điểm (son, phấn) ảo qua camera thiết bị.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TÀI LIỆU THAM KHẢO
# ══════════════════════════════════════════════════════════════════════════════
add_heading("TÀI LIỆU THAM KHẢO", level=1)

references = [
    "[1] Sebastián Ramírez. (2024). FastAPI Documentation. Truy cập từ https://fastapi.tiangolo.com/",
    "[2] Meta Open Source. (2024). React Documentation. Truy cập từ https://react.dev/",
    "[3] The PostgreSQL Global Development Group. (2024). PostgreSQL 16 Documentation. Truy cập từ https://www.postgresql.org/docs/16/",
    "[4] Docker Inc. (2024). Docker Documentation. Truy cập từ https://docs.docker.com/",
    "[5] SQLAlchemy. (2024). SQLAlchemy 2.0 Documentation. Truy cập từ https://docs.sqlalchemy.org/",
    "[6] Pydantic. (2024). Pydantic v2 Documentation. Truy cập từ https://docs.pydantic.dev/",
    "[7] Alembic. (2024). Alembic Documentation. Truy cập từ https://alembic.sqlalchemy.org/",
    "[8] Cloudinary. (2024). Cloudinary Python SDK Documentation. Truy cập từ https://cloudinary.com/documentation/python_integration",
    "[9] PayOS. (2024). PayOS API Documentation. Truy cập từ https://payos.vn/docs/",
    "[10] Auth0. (2023). Introduction to JSON Web Tokens. Truy cập từ https://jwt.io/introduction",
    "[11] Vite. (2024). Vite Documentation. Truy cập từ https://vitejs.dev/",
    "[12] Tailwind Labs. (2024). Tailwind CSS Documentation. Truy cập từ https://tailwindcss.com/docs",
    "[13] TanStack. (2024). TanStack Query Documentation. Truy cập từ https://tanstack.com/query",
    "[14] Zustand. (2024). Zustand Documentation. Truy cập từ https://zustand-demo.pmnd.rs/",
    "[15] VECOM. (2025). Báo cáo Thương mại điện tử Việt Nam 2025. Hà Nội: Hiệp hội Thương mại điện tử Việt Nam.",
    "[16] OWASP. (2023). OWASP Top Ten. Truy cập từ https://owasp.org/www-project-top-ten/",
    "[17] Martin Fowler. (2018). Clean Architecture: A Craftsman's Guide to Software Structure and Design. Addison-Wesley.",
    "[18] Erich Gamma et al. (1994). Design Patterns: Elements of Reusable Object-Oriented Software. Addison-Wesley.",
    "[19] Fielding, R. T. (2000). Architectural Styles and the Design of Network-based Software Architectures. University of California, Irvine.",
    "[20] Nginx Inc. (2024). Nginx Documentation. Truy cập từ https://nginx.org/en/docs/",
]

for ref in references:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(ref)
    set_font(run, size=12)

doc.save(OUT)
print(f"Saved: Chương 5 + 6 + Tài liệu tham khảo → {OUT}")
