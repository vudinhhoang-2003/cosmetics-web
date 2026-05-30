from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.5)
    section.right_margin  = Cm(2.0)

# ── Helper styles ─────────────────────────────────────────────────────────────
def set_font(run, name="Times New Roman", size=13, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rPr.insert(0, rFonts)

def add_paragraph(text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=13, bold=False,
                  italic=False, space_before=0, space_after=6, first_indent=True,
                  color=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(1.0)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic, color=color)
    return p

def add_heading(text, level=1, numbered=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.first_line_indent = Pt(0)
    label = (numbered + " " if numbered else "") + text
    run = p.add_run(label)
    sizes = {1: 16, 2: 14, 3: 13}
    set_font(run, size=sizes.get(level, 13), bold=True)
    return p

def add_bullet(text, indent_level=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(1.0 + indent_level * 0.75)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    run = p.add_run("• " + text)
    set_font(run, size=13)
    return p

def add_table(headers, rows, caption=None):
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_before = Pt(6)
        run = cp.add_run(caption)
        set_font(run, size=12, bold=True, italic=True)

    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # header row
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].add_run(h)
        set_font(run, size=12, bold=True)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'D9E1F2')
        tcPr.append(shd)

    # data rows
    for ri, row_data in enumerate(rows):
        for ci, cell_text in enumerate(row_data):
            cell = tbl.rows[ri + 1].cells[ci]
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = cell.paragraphs[0].add_run(str(cell_text))
            set_font(run, size=12)

    doc.add_paragraph()
    return tbl

# ══════════════════════════════════════════════════════════════════════════════
# TRANG BÌA
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
run = p.add_run("TRƯỜNG ĐẠI HỌC")
set_font(run, size=14, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
run = p.add_run("KHOA CÔNG NGHỆ THÔNG TIN")
set_font(run, size=14, bold=True)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("─────────────────────────────")
set_font(run, size=14)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("BÁO CÁO ĐỒ ÁN MÔN HỌC")
set_font(run, size=18, bold=True)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("WEBSITE THƯƠNG MẠI ĐIỆN TỬ MỸ PHẨM")
set_font(run, size=20, bold=True, color=(31, 78, 121))

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("LUXE BEAUTY")
set_font(run, size=22, bold=True, color=(192, 0, 0))

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Sinh viên thực hiện:  Nguyễn Văn Sơn")
set_font(run, size=14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Giảng viên hướng dẫn:  ...")
set_font(run, size=14)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("TP. Hồ Chí Minh, tháng 5 năm 2026")
set_font(run, size=14, italic=True)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# MỤC LỤC (thủ công)
# ══════════════════════════════════════════════════════════════════════════════
add_heading("MỤC LỤC", level=1)

toc_items = [
    ("CHƯƠNG 1. MỞ ĐẦU", "3"),
    ("    1.1. Lý do chọn đề tài", "3"),
    ("    1.2. Mục tiêu đề tài", "4"),
    ("    1.3. Phạm vi và giới hạn", "5"),
    ("    1.4. Phương pháp thực hiện", "5"),
    ("    1.5. Cấu trúc báo cáo", "6"),
    ("CHƯƠNG 2. CÔNG NGHỆ SỬ DỤNG", "7"),
    ("    2.1. Tổng quan kiến trúc hệ thống", "7"),
    ("    2.2. Backend – FastAPI & Python", "8"),
    ("    2.3. Frontend – React & TypeScript", "12"),
    ("    2.4. Cơ sở dữ liệu – PostgreSQL", "16"),
    ("    2.5. Triển khai – Docker & Nginx", "18"),
    ("    2.6. Tích hợp bên thứ ba", "20"),
    ("CHƯƠNG 3. ĐẶC TẢ YÊU CẦU HỆ THỐNG", "22"),
    ("    3.1. Các bên liên quan", "22"),
    ("    3.2. Yêu cầu chức năng", "23"),
    ("    3.3. Yêu cầu phi chức năng", "30"),
    ("    3.4. Biểu đồ Use Case", "31"),
    ("CHƯƠNG 4. THIẾT KẾ HỆ THỐNG", "33"),
    ("    4.1. Kiến trúc tổng thể", "33"),
    ("    4.2. Thiết kế cơ sở dữ liệu", "35"),
    ("    4.3. Thiết kế API", "39"),
    ("    4.4. Thiết kế giao diện", "45"),
    ("    4.5. Luồng xử lý chính", "47"),
    ("CHƯƠNG 5. CÀI ĐẶT CHƯƠNG TRÌNH VÀ KIỂM THỬ", "50"),
    ("    5.1. Môi trường cài đặt", "50"),
    ("    5.2. Hướng dẫn triển khai", "51"),
    ("    5.3. Kết quả cài đặt – giao diện thực tế", "53"),
    ("    5.4. Kiểm thử hệ thống", "56"),
    ("CHƯƠNG 6. KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", "60"),
    ("    6.1. Kết quả đạt được", "60"),
    ("    6.2. Hạn chế", "61"),
    ("    6.3. Hướng phát triển", "62"),
    ("TÀI LIỆU THAM KHẢO", "63"),
]

for item, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(item)
    set_font(run, size=12, bold=("CHƯƠNG" in item or item.startswith("TÀI")))
    tab_run = p.add_run("\t" + page)
    set_font(tab_run, size=12)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CHƯƠNG 1 – MỞ ĐẦU
# ══════════════════════════════════════════════════════════════════════════════
add_heading("CHƯƠNG 1. MỞ ĐẦU", level=1)

add_heading("1.1. Lý do chọn đề tài", level=2, numbered="1.1.")

add_paragraph(
    "Trong những năm gần đây, thương mại điện tử (TMĐT) tại Việt Nam đã có sự bứt phá "
    "ngoạn mục. Theo báo cáo của Hiệp hội Thương mại điện tử Việt Nam (VECOM), tốc độ "
    "tăng trưởng TMĐT Việt Nam duy trì ở mức trên 25% mỗi năm, đưa Việt Nam trở thành "
    "một trong những thị trường TMĐT phát triển nhanh nhất khu vực Đông Nam Á. Tổng giá "
    "trị giao dịch TMĐT bán lẻ năm 2025 ước đạt hơn 25 tỷ USD, và con số này dự kiến "
    "tiếp tục tăng mạnh trong các năm tới."
)

add_paragraph(
    "Ngành mỹ phẩm – làm đẹp là một trong những lĩnh vực hưởng lợi nhiều nhất từ làn "
    "sóng TMĐT. Người tiêu dùng Việt, đặc biệt là thế hệ Millennials và Gen Z, ngày càng "
    "ưa chuộng việc mua sắm các sản phẩm dưỡng da, trang điểm, nước hoa qua kênh trực "
    "tuyến nhờ sự tiện lợi, đa dạng về mẫu mã, dễ dàng so sánh giá và đọc đánh giá từ "
    "cộng đồng. Quy mô thị trường mỹ phẩm trực tuyến Việt Nam năm 2025 được ước tính vượt "
    "mốc 1,5 tỷ USD, với mức tăng trưởng bình quân trên 20% mỗi năm."
)

add_paragraph(
    "Bên cạnh những ông lớn như Shopee, Lazada hay Tiki, đã có nhiều doanh nghiệp vừa và "
    "nhỏ lựa chọn xây dựng website TMĐT riêng để tạo thương hiệu, chủ động kiểm soát trải "
    "nghiệm khách hàng và không bị phụ thuộc vào chính sách của sàn. Đây là xu hướng "
    "\"Direct-to-Consumer\" (D2C) đang ngày càng phổ biến trên thế giới và tại Việt Nam."
)

add_paragraph(
    "Xuất phát từ thực tiễn đó, đề tài \"Xây dựng Website Thương mại điện tử Mỹ phẩm "
    "LuxeBeauty\" được lựa chọn nhằm nghiên cứu, thiết kế và hiện thực hoá một nền tảng "
    "TMĐT chuyên biệt cho ngành mỹ phẩm. Đề tài không chỉ mang tính thực tiễn cao mà còn "
    "là cơ hội để nhóm sinh viên tổng hợp và áp dụng các kiến thức về lập trình web, "
    "thiết kế cơ sở dữ liệu, bảo mật hệ thống và tích hợp thanh toán vào một sản phẩm "
    "hoàn chỉnh."
)

add_paragraph(
    "Hơn nữa, việc sử dụng các công nghệ hiện đại như FastAPI, React, PostgreSQL và Docker "
    "trong dự án này giúp sinh viên làm quen với quy trình phát triển phần mềm chuyên "
    "nghiệp, tiếp cận kiến trúc microservice và nắm vững kỹ năng triển khai ứng dụng "
    "trong môi trường container hoá – những kỹ năng đang được thị trường lao động công "
    "nghệ thông tin đòi hỏi cao."
)

add_heading("1.2. Mục tiêu đề tài", level=2, numbered="1.2.")

add_paragraph("Đề tài hướng đến các mục tiêu cụ thể sau đây:")

add_bullet("Xây dựng một website thương mại điện tử mỹ phẩm đầy đủ chức năng, bao gồm "
           "trang hiển thị sản phẩm, giỏ hàng, thanh toán, quản lý đơn hàng và hệ thống "
           "đánh giá sản phẩm.")

add_bullet("Thiết kế và triển khai hệ thống xác thực người dùng (đăng ký, đăng nhập) "
           "sử dụng JWT (JSON Web Token) với phân quyền theo vai trò (khách hàng / quản trị viên).")

add_bullet("Xây dựng trang quản trị (Admin Dashboard) cho phép quản lý toàn bộ sản phẩm, "
           "danh mục, đơn hàng, người dùng và xem thống kê doanh thu.")

add_bullet("Tích hợp cổng thanh toán trực tuyến PayOS – một giải pháp thanh toán phổ biến "
           "tại Việt Nam, hỗ trợ QR Code và chuyển khoản ngân hàng.")

add_bullet("Áp dụng kiến trúc phân lớp rõ ràng (Clean Architecture) cho backend: "
           "Models → Schemas → CRUD → Routers, đảm bảo tính bảo trì và mở rộng.")

add_bullet("Đóng gói toàn bộ hệ thống bằng Docker Compose, đảm bảo môi trường triển khai "
           "nhất quán giữa development và production.")

add_bullet("Viết bộ kiểm thử tự động (automated tests) cho các API quan trọng bằng Pytest, "
           "đảm bảo chất lượng mã nguồn.")

add_bullet("Nghiên cứu và trình bày các công nghệ, thư viện được sử dụng, làm nền tảng lý "
           "thuyết vững chắc cho việc hiện thực hoá dự án.")

add_heading("1.3. Phạm vi và giới hạn", level=2, numbered="1.3.")

add_paragraph(
    "Dự án tập trung vào việc xây dựng một website TMĐT mỹ phẩm hoàn chỉnh với các chức "
    "năng cốt lõi. Phạm vi thực hiện bao gồm:"
)

add_bullet("Phân hệ khách hàng: Duyệt và tìm kiếm sản phẩm, quản lý giỏ hàng, đặt hàng, "
           "thanh toán, theo dõi đơn hàng, đánh giá sản phẩm, quản lý hồ sơ cá nhân.")

add_bullet("Phân hệ quản trị: Quản lý sản phẩm (CRUD), quản lý danh mục, quản lý đơn "
           "hàng (cập nhật trạng thái), quản lý người dùng, xem thống kê tổng quan.")

add_bullet("Hệ thống backend RESTful API xây dựng bằng FastAPI, cơ sở dữ liệu PostgreSQL, "
           "lưu trữ ảnh trên Cloudinary.")

add_bullet("Giao diện frontend xây dựng bằng React + TypeScript + Tailwind CSS, "
           "responsive trên các thiết bị.")

add_paragraph("Các giới hạn của dự án:")

add_bullet("Chưa tích hợp hệ thống giao hàng thực tế (chỉ lưu địa chỉ giao hàng).")
add_bullet("Chưa có tính năng chat trực tuyến, hỗ trợ khách hàng realtime.")
add_bullet("Chưa tích hợp đầy đủ các phương thức thanh toán quốc tế (Visa, Mastercard).")
add_bullet("Chức năng mã giảm giá (coupon/voucher) chưa được triển khai trong phiên bản này.")
add_bullet("Ứng dụng di động (iOS/Android) nằm ngoài phạm vi của đề tài này.")

add_heading("1.4. Phương pháp thực hiện", level=2, numbered="1.4.")

add_paragraph(
    "Để hoàn thành đề tài, các phương pháp nghiên cứu và phát triển sau được áp dụng:"
)

add_bullet("Nghiên cứu tài liệu: Tìm hiểu tài liệu chính thức của FastAPI, React, "
           "PostgreSQL, Docker, JWT, Cloudinary và PayOS API. Tham khảo các bài viết kỹ "
           "thuật, khóa học trực tuyến và mã nguồn mở trên GitHub.")

add_bullet("Phân tích hệ thống tương tự: Khảo sát các website TMĐT mỹ phẩm trong nước "
           "và quốc tế (Hasaki.vn, Guardian, Sephora) để tham khảo quy trình người dùng "
           "và các tính năng cốt lõi cần thiết.")

add_bullet("Phát triển theo phương pháp Agile: Chia dự án thành các sprint ngắn, mỗi "
           "sprint tập trung vào một nhóm tính năng. Sử dụng Git để quản lý phiên bản mã nguồn.")

add_bullet("Kiểm thử liên tục: Viết unit test và integration test bằng Pytest cho backend. "
           "Kiểm tra thủ công giao diện trên nhiều trình duyệt và kích thước màn hình.")

add_bullet("Tổng hợp và viết báo cáo: Ghi chép lại quá trình thiết kế, quyết định kỹ thuật "
           "và kết quả đạt được trong báo cáo này.")

add_heading("1.5. Cấu trúc báo cáo", level=2, numbered="1.5.")

add_paragraph("Báo cáo được tổ chức thành 6 chương với nội dung cụ thể như sau:")

add_bullet("Chương 1 – Mở đầu: Trình bày lý do chọn đề tài, mục tiêu, phạm vi và phương "
           "pháp thực hiện.")

add_bullet("Chương 2 – Công nghệ sử dụng: Giới thiệu chi tiết các công nghệ, framework, "
           "thư viện được áp dụng trong dự án cùng lý do lựa chọn.")

add_bullet("Chương 3 – Đặc tả yêu cầu hệ thống: Mô tả các yêu cầu chức năng và phi chức "
           "năng, biểu đồ Use Case thể hiện tương tác của người dùng với hệ thống.")

add_bullet("Chương 4 – Thiết kế hệ thống: Trình bày kiến trúc tổng thể, thiết kế cơ sở "
           "dữ liệu, thiết kế API và giao diện người dùng.")

add_bullet("Chương 5 – Cài đặt chương trình và kiểm thử: Hướng dẫn triển khai, trình bày "
           "kết quả cài đặt thực tế và kết quả kiểm thử hệ thống.")

add_bullet("Chương 6 – Kết luận và hướng phát triển: Tổng kết kết quả đạt được, chỉ ra "
           "hạn chế và đề xuất hướng phát triển trong tương lai.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CHƯƠNG 2 – CÔNG NGHỆ SỬ DỤNG
# ══════════════════════════════════════════════════════════════════════════════
add_heading("CHƯƠNG 2. CÔNG NGHỆ SỬ DỤNG", level=1)

add_heading("2.1. Tổng quan kiến trúc hệ thống", level=2, numbered="2.1.")

add_paragraph(
    "Hệ thống LuxeBeauty được xây dựng theo kiến trúc phân lớp ba tầng (3-tier architecture) "
    "kết hợp mô hình Client-Server truyền thống với các thực hành tốt nhất của ứng dụng "
    "web hiện đại. Toàn bộ hệ thống được đóng gói và triển khai bằng Docker Compose, "
    "đảm bảo tính nhất quán môi trường và khả năng mở rộng."
)

add_paragraph("Kiến trúc hệ thống bao gồm các lớp chính:")

add_bullet("Tầng trình diện (Presentation Layer): Ứng dụng React (Single Page Application) "
           "chạy trên trình duyệt, giao tiếp với backend thông qua RESTful API.")

add_bullet("Tầng xử lý nghiệp vụ (Business Logic Layer): FastAPI server xử lý toàn bộ "
           "logic nghiệp vụ, xác thực, phân quyền và điều phối các tác vụ.")

add_bullet("Tầng dữ liệu (Data Layer): PostgreSQL lưu trữ dữ liệu quan hệ, Cloudinary "
           "lưu trữ hình ảnh sản phẩm trên đám mây.")

add_bullet("Lớp Proxy (Reverse Proxy): Nginx đóng vai trò cổng vào (gateway) điều hướng "
           "request từ client đến đúng service.")

add_paragraph(
    "Luồng xử lý tổng quát: Người dùng truy cập website qua trình duyệt → Nginx nhận "
    "request và phân loại → Request đến API (tiền tố /api/) được chuyển đến FastAPI backend "
    "→ Backend truy vấn PostgreSQL và trả về dữ liệu JSON → Frontend React nhận dữ liệu "
    "và render giao diện. Ảnh sản phẩm được lưu trực tiếp lên Cloudinary CDN và phân "
    "phối đến người dùng qua URL Cloudinary."
)

add_table(
    ["Thành phần", "Công nghệ", "Phiên bản", "Vai trò"],
    [
        ["Backend", "FastAPI + Python", "0.111.0 / 3.12", "API server, xử lý nghiệp vụ"],
        ["Frontend", "React + TypeScript", "18.3.1 / 5.4.5", "Giao diện người dùng (SPA)"],
        ["Database", "PostgreSQL", "16", "Lưu trữ dữ liệu quan hệ"],
        ["ORM", "SQLAlchemy", "2.0.30", "Ánh xạ đối tượng – cơ sở dữ liệu"],
        ["Migration", "Alembic", "1.13.1", "Quản lý phiên bản schema DB"],
        ["Reverse Proxy", "Nginx", "Alpine", "Cổng vào, phân phối tĩnh"],
        ["Container", "Docker + Compose", "Latest", "Đóng gói và triển khai"],
        ["Image CDN", "Cloudinary", "1.40.0", "Lưu trữ và phân phối ảnh"],
        ["Payment", "PayOS", "Latest", "Cổng thanh toán trực tuyến"],
        ["Build Tool", "Vite", "5.3.1", "Build frontend nhanh"],
    ],
    caption="Bảng 2.1. Tổng hợp công nghệ sử dụng trong dự án"
)

add_heading("2.2. Backend – FastAPI & Python", level=2, numbered="2.2.")

add_heading("2.2.1. Giới thiệu FastAPI", level=3)

add_paragraph(
    "FastAPI là một web framework hiện đại, hiệu năng cao được xây dựng trên Python 3.7+ "
    "và dựa trên tiêu chuẩn OpenAPI (Swagger) cùng JSON Schema. Framework này được tạo "
    "ra bởi Sebastián Ramírez và lần đầu ra mắt vào năm 2018. FastAPI nhanh chóng trở "
    "thành một trong những framework Python phổ biến nhất nhờ các ưu điểm vượt trội."
)

add_paragraph(
    "FastAPI được thiết kế để đạt hiệu năng cao nhờ sử dụng Starlette làm nền tảng ASGI "
    "(Asynchronous Server Gateway Interface) và Pydantic để validation dữ liệu. Theo "
    "các benchmark độc lập, FastAPI đạt hiệu năng ngang bằng với NodeJS và Go, cao hơn "
    "nhiều so với Django REST Framework hay Flask. Điều này đặc biệt quan trọng cho các "
    "ứng dụng TMĐT có lưu lượng truy cập lớn."
)

add_paragraph("Các ưu điểm chính của FastAPI được sử dụng trong dự án:")

add_bullet("Tự động sinh tài liệu API (Swagger UI tại /docs, ReDoc tại /redoc) từ code, "
           "giúp frontend dễ dàng tích hợp và kiểm thử API.")
add_bullet("Validation tự động dữ liệu request/response thông qua Pydantic v2, giảm thiểu "
           "lỗi runtime và bảo mật hơn.")
add_bullet("Hỗ trợ async/await natively, cho phép xử lý nhiều request đồng thời mà không "
           "block thread.")
add_bullet("Dependency Injection tích hợp sẵn, giúp quản lý các phụ thuộc (database session, "
           "current user) một cách sạch sẽ.")
add_bullet("Type hints Python được tận dụng triệt để, giúp IDE hỗ trợ tốt và giảm bug.")

add_heading("2.2.2. Pydantic v2 – Validation dữ liệu", level=3)

add_paragraph(
    "Pydantic là thư viện validation dữ liệu mạnh mẽ nhất trong hệ sinh thái Python. "
    "Phiên bản 2 (Pydantic v2) được viết lại hoàn toàn bằng Rust, mang lại tốc độ "
    "nhanh hơn 5-50 lần so với v1. Trong dự án LuxeBeauty, Pydantic được sử dụng để:"
)

add_bullet("Định nghĩa các schema cho request body (ví dụ: ProductCreate, UserCreate, OrderCreate).")
add_bullet("Định nghĩa các schema cho response (ví dụ: ProductOut, UserOut, OrderOut) với "
           "chỉ những trường cần thiết.")
add_bullet("Validate dữ liệu đầu vào tự động: kiểu dữ liệu, giá trị tối thiểu/tối đa, "
           "định dạng email, độ dài chuỗi.")
add_bullet("Sử dụng model_config với from_attributes=True để chuyển đổi từ SQLAlchemy ORM "
           "object sang Pydantic schema.")

add_heading("2.2.3. SQLAlchemy & Alembic", level=3)

add_paragraph(
    "SQLAlchemy là ORM (Object-Relational Mapper) phổ biến nhất trong Python. Phiên bản "
    "2.0 mang lại cú pháp mới gọn hơn với mapped_column() và Mapped[] type annotation. "
    "Trong dự án, SQLAlchemy được sử dụng để:"
)

add_bullet("Định nghĩa các model ORM ánh xạ trực tiếp với bảng trong PostgreSQL "
           "(User, Product, Category, Order, CartItem, Review).")
add_bullet("Thực hiện truy vấn dữ liệu thông qua Session API với các phương thức như "
           "select(), scalars(), execute().")
add_bullet("Quản lý transaction tự động thông qua context manager.")
add_bullet("Hỗ trợ mối quan hệ giữa các bảng (relationship()) với lazy loading.")

add_paragraph(
    "Alembic là công cụ migration cơ sở dữ liệu chính thức của SQLAlchemy. Dự án sử dụng "
    "Alembic để quản lý lịch sử thay đổi schema theo từng phiên bản, đảm bảo việc cập "
    "nhật cấu trúc database không làm mất dữ liệu. Dự án có 2 migration: khởi tạo "
    "schema ban đầu và thêm trường order_code."
)

add_heading("2.2.4. Bảo mật – JWT & Bcrypt", level=3)

add_paragraph(
    "JSON Web Token (JWT) là tiêu chuẩn mở (RFC 7519) để truyền tải thông tin an toàn "
    "giữa các bên dưới dạng JSON object. JWT được ký bằng thuật toán HS256 (HMAC với "
    "SHA-256) sử dụng Secret Key. Hệ thống xác thực của LuxeBeauty gồm:"
)

add_bullet("Access Token: Thời hạn ngắn (30 phút mặc định), dùng để xác thực mỗi request.")
add_bullet("Refresh Token: Thời hạn dài hơn (7 ngày), dùng để cấp lại Access Token khi hết hạn.")
add_bullet("HTTPBearer scheme: Token được gửi trong Authorization header với định dạng "
           "\"Bearer {token}\".")
add_bullet("Dependency get_current_user: Tự động giải mã và xác thực token trong mỗi "
           "endpoint yêu cầu xác thực.")
add_bullet("Dependency require_admin: Kiểm tra vai trò admin trước khi cho phép truy cập "
           "các endpoint quản trị.")

add_paragraph(
    "Bcrypt là thuật toán hash mật khẩu mạnh, thiết kế chống lại brute-force attack nhờ "
    "hệ số cost có thể điều chỉnh. Thư viện bcrypt==4.2.1 được sử dụng để hash mật khẩu "
    "khi đăng ký và verify khi đăng nhập. Mật khẩu không bao giờ được lưu dưới dạng "
    "plaintext trong database."
)

add_heading("2.2.5. Rate Limiting – SlowAPI", level=3)

add_paragraph(
    "SlowAPI là thư viện rate limiting cho FastAPI, xây dựng trên nền Limits. Rate limiting "
    "giúp bảo vệ API khỏi các cuộc tấn công DDoS và abuse. Trong dự án, rate limiting "
    "được áp dụng tại các endpoint nhạy cảm như đăng nhập và đăng ký, giới hạn số lần "
    "gọi API trong một khoảng thời gian nhất định từ mỗi địa chỉ IP."
)

add_heading("2.3. Frontend – React & TypeScript", level=2, numbered="2.3.")

add_heading("2.3.1. React 18", level=3)

add_paragraph(
    "React là thư viện JavaScript phổ biến nhất thế giới để xây dựng giao diện người dùng, "
    "được phát triển bởi Meta (Facebook) từ năm 2013. React 18 (phát hành 2022) mang đến "
    "nhiều cải tiến quan trọng:"
)

add_bullet("Concurrent Rendering: Cho phép React xử lý nhiều tác vụ UI đồng thời, cải "
           "thiện độ mượt mà của giao diện.")
add_bullet("Automatic Batching: Tự động gom nhóm nhiều setState calls thành một lần render, "
           "tăng hiệu năng.")
add_bullet("Suspense cho Data Fetching: Cải thiện trải nghiệm loading state.")
add_bullet("useId Hook mới: Tạo ID duy nhất phục vụ accessibility.")

add_paragraph(
    "Dự án LuxeBeauty sử dụng React theo mô hình Component-Based Architecture, nơi giao "
    "diện được chia nhỏ thành các component tái sử dụng: Layout, Navbar, Footer, "
    "ProductCard, v.v. Routing được xử lý bởi React Router DOM v6."
)

add_heading("2.3.2. TypeScript", level=3)

add_paragraph(
    "TypeScript là ngôn ngữ lập trình được Microsoft phát triển, thêm static typing vào "
    "JavaScript. TypeScript compiles thành JavaScript thuần, có thể chạy trên mọi môi "
    "trường. Lợi ích chính trong dự án:"
)

add_bullet("Type Safety: Phát hiện lỗi kiểu dữ liệu ở compile time thay vì runtime, "
           "giảm bug liên quan đến null/undefined và kiểu dữ liệu sai.")
add_bullet("IntelliSense: IDE (VSCode) hỗ trợ autocomplete, refactor, và navigation mã nguồn "
           "tốt hơn nhiều so với JavaScript.")
add_bullet("Interface và Type: Định nghĩa rõ ràng cấu trúc dữ liệu (User, Product, Order) "
           "trong file types/index.ts, đảm bảo nhất quán giữa các component.")
add_bullet("Tích hợp tốt với React: TypeScript + React = type-safe props, events và hooks.")

add_heading("2.3.3. Vite – Build Tool", level=3)

add_paragraph(
    "Vite (phát âm \"vít\") là build tool thế hệ mới được Evan You (tác giả Vue.js) "
    "phát triển. Vite giải quyết điểm yếu lớn nhất của Webpack: tốc độ khởi động slow "
    "và Hot Module Replacement (HMR) chậm khi dự án lớn."
)

add_bullet("Dev Server: Sử dụng ESModules natively trong trình duyệt, không bundle khi "
           "development → khởi động gần như tức thì.")
add_bullet("HMR cực nhanh: Chỉ cập nhật đúng module thay đổi, không reload toàn bộ trang.")
add_bullet("Build Production: Sử dụng Rollup để bundle tối ưu cho production.")
add_bullet("Plugin Ecosystem: Tích hợp tốt với React, TypeScript, Tailwind CSS.")
add_bullet("Proxy Configuration: Vite config hỗ trợ proxy /api/* đến backend trong dev, "
           "tránh CORS issues.")

add_heading("2.3.4. Tailwind CSS", level=3)

add_paragraph(
    "Tailwind CSS là framework CSS utility-first, thay vì viết CSS tùy chỉnh, developer "
    "sử dụng các class tiện ích được định sẵn (flex, p-4, text-center, bg-pink-500...) "
    "trực tiếp trong HTML/JSX. Đây là sự khác biệt lớn so với Bootstrap (component-based)."
)

add_bullet("Không cần đặt tên class: Loại bỏ vấn đề đặt tên CSS class và xung đột style.")
add_bullet("PurgeCSS tích hợp: Tự động loại bỏ các utility class không dùng khi build, "
           "file CSS production rất nhỏ gọn.")
add_bullet("Responsive Design: Tiền tố md:, lg:, xl: giúp xây dựng responsive layout "
           "nhanh chóng.")
add_bullet("Dark Mode: Hỗ trợ dark mode với tiền tố dark: (có thể mở rộng trong tương lai).")
add_bullet("Tùy biến cao: tailwind.config.js cho phép mở rộng màu sắc, font, spacing theo "
           "brand guideline của LuxeBeauty.")

add_heading("2.3.5. Zustand – Quản lý trạng thái", level=3)

add_paragraph(
    "Zustand là thư viện quản lý trạng thái (state management) nhẹ nhàng cho React, "
    "được Bear (Poimandres) phát triển. So với Redux, Zustand có API đơn giản hơn nhiều "
    "và ít boilerplate code hơn. Dự án sử dụng 3 Zustand store:"
)

add_bullet("authStore: Lưu thông tin người dùng đăng nhập (user object, access token, "
           "refresh token), trạng thái isAuthenticated. Hỗ trợ login/logout action.")
add_bullet("adminAuthStore: Store riêng biệt cho xác thực admin, tách biệt với customer "
           "session để tránh xung đột.")
add_bullet("cartStore: Quản lý giỏ hàng phía client (danh sách sản phẩm, số lượng, tổng "
           "tiền). Đồng bộ với server khi user đăng nhập.")

add_heading("2.3.6. TanStack Query (React Query)", level=3)

add_paragraph(
    "TanStack Query (trước đây là React Query) là thư viện mạnh mẽ để quản lý server "
    "state trong React. Thư viện này giải quyết các vấn đề phức tạp như caching, "
    "background refetching, optimistic updates. Trong dự án:"
)

add_bullet("useQuery: Fetch và cache dữ liệu sản phẩm, danh mục, đơn hàng với stale time "
           "phù hợp, tránh gọi API không cần thiết.")
add_bullet("useMutation: Xử lý các tác vụ POST/PUT/DELETE với error handling và "
           "invalidation cache tự động.")
add_bullet("Automatic retry: Tự động retry khi request thất bại (mạng yếu).")

add_heading("2.3.7. Framer Motion – Animation", level=3)

add_paragraph(
    "Framer Motion là thư viện animation phổ biến nhất cho React, được dùng tại LuxeBeauty "
    "để tạo các hiệu ứng chuyển động mượt mà: fade-in khi trang load, slide-up cho "
    "product cards, hover effects cho nút bấm. Điều này cải thiện đáng kể trải nghiệm "
    "người dùng và tạo cảm giác cao cấp phù hợp với brand mỹ phẩm."
)

add_heading("2.4. Cơ sở dữ liệu – PostgreSQL", level=2, numbered="2.4.")

add_paragraph(
    "PostgreSQL (thường gọi là Postgres) là hệ quản trị cơ sở dữ liệu quan hệ mã nguồn "
    "mở mạnh mẽ nhất hiện nay, với lịch sử phát triển hơn 35 năm. PostgreSQL nổi bật với "
    "sự tuân thủ chặt chẽ tiêu chuẩn SQL, tính toàn vẹn dữ liệu cao và khả năng mở rộng tốt."
)

add_paragraph("Lý do chọn PostgreSQL cho dự án LuxeBeauty:")

add_bullet("ACID Compliance: Đảm bảo tính toàn vẹn giao dịch hoàn toàn, quan trọng cho "
           "các thao tác đặt hàng và thanh toán.")
add_bullet("JSON/JSONB Support: Cột shipping_address trong bảng orders sử dụng JSONB để "
           "lưu địa chỉ linh hoạt mà không cần normalize.")
add_bullet("Array Type: Cột images trong bảng products sử dụng kiểu ARRAY(Text) của "
           "PostgreSQL để lưu nhiều URL ảnh.")
add_bullet("UUID Support: Hỗ trợ kiểu dữ liệu UUID natively thông qua extension, tốt "
           "hơn auto-increment integer cho distributed systems.")
add_bullet("Concurrent Access: Xử lý tốt nhiều kết nối đồng thời, phù hợp với website "
           "TMĐT có traffic cao.")
add_bullet("Rich Query Language: Window functions, CTEs, Full-text search hỗ trợ các "
           "truy vấn phức tạp cho báo cáo doanh thu.")

add_paragraph(
    "Phiên bản PostgreSQL 16 được sử dụng, chạy trong Docker container với volume "
    "pgdata để đảm bảo dữ liệu không mất khi restart container."
)

add_heading("2.5. Triển khai – Docker & Nginx", level=2, numbered="2.5.")

add_heading("2.5.1. Docker & Docker Compose", level=3)

add_paragraph(
    "Docker là nền tảng containerization cho phép đóng gói ứng dụng cùng toàn bộ "
    "dependencies vào một container độc lập. Container đảm bảo ứng dụng chạy nhất quán "
    "trên mọi môi trường (development laptop, staging server, production cloud)."
)

add_paragraph(
    "Docker Compose là công cụ định nghĩa và chạy ứng dụng multi-container. File "
    "docker-compose.yml của dự án định nghĩa 3 service chính:"
)

add_bullet("db: PostgreSQL 16 Alpine – nhẹ nhất, healthcheck đảm bảo DB sẵn sàng trước "
           "khi backend khởi động.")
add_bullet("backend: FastAPI + Uvicorn – depends_on db với điều kiện service_healthy, "
           "volume uploads cho ảnh upload local.")
add_bullet("frontend: React build được serve bởi Nginx Alpine – multi-stage Dockerfile "
           "giúp image production nhỏ gọn (< 50MB).")

add_paragraph(
    "Multi-stage build trong Dockerfile frontend: Stage 1 (node:20-alpine) cài dependencies "
    "và build React app → Stage 2 (nginx:alpine) chỉ copy dist/ folder, không mang theo "
    "node_modules (~300MB). Kết quả: image production nhỏ và an toàn hơn."
)

add_heading("2.5.2. Nginx – Reverse Proxy", level=3)

add_paragraph(
    "Nginx (đọc là \"engine-x\") là web server và reverse proxy hiệu năng cao. Trong dự "
    "án, Nginx đóng vai trò:"
)

add_bullet("Reverse Proxy: Điều hướng request /api/* đến FastAPI backend:8000, các request "
           "còn lại đến frontend.")
add_bullet("Static File Server: Serve các file tĩnh (HTML, CSS, JS) của React app với "
           "caching headers tối ưu.")
add_bullet("SPA Support: Cấu hình try_files $uri $uri/ /index.html để hỗ trợ React Router "
           "– mọi path đều trả về index.html.")
add_bullet("Upload Size: client_max_body_size 10M cho phép upload ảnh lên đến 10MB.")
add_bullet("Compression: gzip compression giảm bandwidth và tăng tốc độ tải trang.")

add_heading("2.6. Tích hợp bên thứ ba", level=2, numbered="2.6.")

add_heading("2.6.1. Cloudinary – Lưu trữ ảnh", level=3)

add_paragraph(
    "Cloudinary là nền tảng quản lý media (ảnh, video) trên đám mây, được hơn 1 triệu "
    "developer và doanh nghiệp sử dụng. Cloudinary cung cấp CDN toàn cầu để phân phối "
    "ảnh nhanh nhất đến người dùng. Trong dự án:"
)

add_bullet("Upload API: Endpoint POST /api/upload/image nhận file ảnh từ admin, upload "
           "lên Cloudinary và trả về URL công khai.")
add_bullet("Transformation: Cloudinary hỗ trợ resize, crop, optimize ảnh on-the-fly "
           "thông qua URL parameters (chưa khai thác đầy đủ trong phiên bản này).")
add_bullet("Auto format: Cloudinary tự động chuyển ảnh sang WebP khi trình duyệt hỗ trợ, "
           "giảm kích thước file đáng kể.")

add_heading("2.6.2. PayOS – Cổng thanh toán", level=3)

add_paragraph(
    "PayOS là cổng thanh toán trực tuyến của Việt Nam, hỗ trợ thanh toán qua QR Code "
    "ngân hàng và chuyển khoản liên ngân hàng. PayOS được nhiều startup và SME Việt Nam "
    "sử dụng nhờ tích hợp đơn giản và phí thấp."
)

add_bullet("Tạo link thanh toán: API tạo payment link với thông tin đơn hàng, số tiền, "
           "order code.")
add_bullet("Webhook: PayOS gửi webhook đến /api/payment/webhook khi giao dịch thành công, "
           "backend cập nhật trạng thái đơn hàng.")
add_bullet("Mock Payment: Endpoint /api/payment/simulate-success cho phép test luồng "
           "thanh toán mà không cần giao dịch thật.")
add_bullet("COD: Ngoài PayOS, hệ thống hỗ trợ hình thức thanh toán khi nhận hàng (Cash "
           "on Delivery).")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# LƯU FILE SAU CHƯƠNG 1+2
# ══════════════════════════════════════════════════════════════════════════════
doc.save(r"f:\DOANTHUE\Son\Website_Cosmetics\BaoCao_LuxeBeauty.docx")
print("Saved: Chương 1 + 2 xong")
