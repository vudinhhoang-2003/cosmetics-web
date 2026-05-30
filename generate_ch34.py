from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document(r"f:\DOANTHUE\Son\Website_Cosmetics\BaoCao_LuxeBeauty.docx")
OUT = r"f:\DOANTHUE\Son\Website_Cosmetics\BaoCao_LuxeBeauty_tmp.docx"

def set_font(run, name="Times New Roman", size=13, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rPr.insert(0, rFonts)

def add_paragraph(text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=13, bold=False,
                  italic=False, space_before=0, space_after=6, first_indent=True):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(1.0)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic)
    return p

def add_heading(text, level=1, numbered=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.first_line_indent = Pt(0)
    label = (numbered + " " if numbered else "") + text
    run = p.add_run(label)
    sizes = {1: 16, 2: 14, 3: 13}
    set_font(run, size=sizes.get(level, 13), bold=True)
    return p

def add_bullet(text, indent_level=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(1.0 + indent_level * 0.75)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    run = p.add_run("• " + text)
    set_font(run, size=13)
    return p

def add_table(headers, rows, caption=None):
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_before = Pt(6)
        cp.paragraph_format.first_line_indent = Pt(0)
        run = cp.add_run(caption)
        set_font(run, size=12, bold=True, italic=True)

    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].add_run(h)
        set_font(run, size=12, bold=True)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'D9E1F2')
        tcPr.append(shd)

    for ri, row_data in enumerate(rows):
        for ci, cell_text in enumerate(row_data):
            cell = tbl.rows[ri + 1].cells[ci]
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = cell.paragraphs[0].add_run(str(cell_text))
            set_font(run, size=12)

    doc.add_paragraph()
    return tbl

# ══════════════════════════════════════════════════════════════════════════════
# CHƯƠNG 3 – ĐẶC TẢ YÊU CẦU HỆ THỐNG
# ══════════════════════════════════════════════════════════════════════════════
add_heading("CHƯƠNG 3. ĐẶC TẢ YÊU CẦU HỆ THỐNG", level=1)

add_heading("3.1. Các bên liên quan (Stakeholders)", level=2, numbered="3.1.")

add_paragraph(
    "Hệ thống LuxeBeauty phục vụ ba nhóm người dùng chính với nhu cầu và quyền hạn khác nhau:"
)

add_table(
    ["Bên liên quan", "Mô tả", "Mục tiêu chính"],
    [
        ["Khách hàng (Customer)",
         "Người dùng cuối mua sắm trực tuyến, có thể là khách vãng lai hoặc đã đăng ký tài khoản",
         "Tìm kiếm, mua sản phẩm mỹ phẩm nhanh chóng, tiện lợi"],
        ["Quản trị viên (Admin)",
         "Nhân viên hoặc chủ cửa hàng quản lý toàn bộ hệ thống",
         "Quản lý sản phẩm, đơn hàng, người dùng; xem báo cáo kinh doanh"],
        ["Hệ thống bên ngoài",
         "PayOS (cổng thanh toán), Cloudinary (lưu trữ ảnh)",
         "Cung cấp dịch vụ thanh toán và lưu trữ media"],
    ],
    caption="Bảng 3.1. Các bên liên quan trong hệ thống"
)

add_heading("3.2. Yêu cầu chức năng", level=2, numbered="3.2.")

add_heading("3.2.1. Phân hệ Xác thực (Authentication)", level=3)

add_paragraph("Bảng dưới đây mô tả các yêu cầu chức năng cho phân hệ xác thực:")

add_table(
    ["Mã YC", "Tên chức năng", "Mô tả", "Đối tượng"],
    [
        ["UC-01", "Đăng ký tài khoản",
         "Người dùng nhập họ tên, email, mật khẩu (tối thiểu 8 ký tự, có chữ hoa, số) để tạo tài khoản mới. Hệ thống kiểm tra email chưa tồn tại.",
         "Khách vãng lai"],
        ["UC-02", "Đăng nhập",
         "Người dùng nhập email và mật khẩu. Hệ thống xác thực và cấp Access Token + Refresh Token.",
         "Khách hàng, Admin"],
        ["UC-03", "Làm mới token",
         "Khi Access Token hết hạn, hệ thống tự động sử dụng Refresh Token để cấp token mới mà không cần đăng nhập lại.",
         "Khách hàng, Admin"],
        ["UC-04", "Đăng xuất",
         "Xoá token khỏi client-side storage, kết thúc phiên làm việc.",
         "Khách hàng, Admin"],
        ["UC-05", "Phân quyền theo vai trò",
         "Hệ thống phân biệt vai trò 'customer' và 'admin'. Admin có quyền truy cập các endpoint quản trị, customer bị từ chối.",
         "Hệ thống"],
    ],
    caption="Bảng 3.2. Yêu cầu chức năng phân hệ xác thực"
)

add_heading("3.2.2. Phân hệ Sản phẩm (Product)", level=3)

add_table(
    ["Mã YC", "Tên chức năng", "Mô tả", "Đối tượng"],
    [
        ["UC-06", "Xem danh sách sản phẩm",
         "Hiển thị danh sách sản phẩm có phân trang (20 sản phẩm/trang). Hỗ trợ lọc theo danh mục, khoảng giá, thương hiệu, trạng thái tồn kho.",
         "Tất cả"],
        ["UC-07", "Tìm kiếm sản phẩm",
         "Tìm kiếm sản phẩm theo tên hoặc mô tả (case-insensitive, ILIKE). Kết quả cập nhật realtime khi nhập.",
         "Tất cả"],
        ["UC-08", "Sắp xếp sản phẩm",
         "Sắp xếp theo giá tăng/giảm dần, theo mức độ mới nhất, theo lượt đánh giá.",
         "Tất cả"],
        ["UC-09", "Xem chi tiết sản phẩm",
         "Hiển thị thông tin đầy đủ: tên, mô tả, giá, giá khuyến mãi, hình ảnh (carousel), thương hiệu, tồn kho, danh mục, đánh giá.",
         "Tất cả"],
        ["UC-10", "Thêm sản phẩm (Admin)",
         "Admin thêm sản phẩm mới với đầy đủ thông tin. Upload nhiều ảnh lên Cloudinary. Tự động tạo slug từ tên.",
         "Admin"],
        ["UC-11", "Sửa sản phẩm (Admin)",
         "Cập nhật thông tin sản phẩm. Có thể thêm/xoá ảnh.",
         "Admin"],
        ["UC-12", "Xoá sản phẩm (Admin)",
         "Xoá mềm (set is_active = false) hoặc xoá cứng sản phẩm khỏi hệ thống.",
         "Admin"],
    ],
    caption="Bảng 3.3. Yêu cầu chức năng phân hệ sản phẩm"
)

add_heading("3.2.3. Phân hệ Giỏ hàng (Cart)", level=3)

add_table(
    ["Mã YC", "Tên chức năng", "Mô tả", "Đối tượng"],
    [
        ["UC-13", "Xem giỏ hàng",
         "Hiển thị danh sách sản phẩm trong giỏ hàng kèm số lượng, đơn giá, thành tiền và tổng tiền.",
         "Khách hàng đã đăng nhập"],
        ["UC-14", "Thêm vào giỏ hàng",
         "Thêm sản phẩm vào giỏ. Nếu sản phẩm đã có, tăng số lượng. Kiểm tra tồn kho trước khi thêm.",
         "Khách hàng đã đăng nhập"],
        ["UC-15", "Cập nhật số lượng",
         "Thay đổi số lượng sản phẩm trong giỏ. Số lượng tối thiểu là 1.",
         "Khách hàng đã đăng nhập"],
        ["UC-16", "Xoá khỏi giỏ hàng",
         "Xoá một hoặc tất cả sản phẩm khỏi giỏ hàng.",
         "Khách hàng đã đăng nhập"],
    ],
    caption="Bảng 3.4. Yêu cầu chức năng phân hệ giỏ hàng"
)

add_heading("3.2.4. Phân hệ Đặt hàng & Thanh toán", level=3)

add_table(
    ["Mã YC", "Tên chức năng", "Mô tả", "Đối tượng"],
    [
        ["UC-17", "Đặt hàng",
         "Tạo đơn hàng từ giỏ hàng hiện tại. Nhập địa chỉ giao hàng (tên, số điện thoại, địa chỉ, tỉnh/thành). Chọn phương thức thanh toán.",
         "Khách hàng đã đăng nhập"],
        ["UC-18", "Thanh toán online (PayOS)",
         "Tạo payment link qua PayOS API, chuyển hướng đến trang thanh toán. Sau khi thanh toán thành công, PayOS gửi webhook xác nhận.",
         "Khách hàng đã đăng nhập"],
        ["UC-19", "Thanh toán COD",
         "Chọn thanh toán khi nhận hàng. Đơn hàng được tạo với trạng thái 'pending'.",
         "Khách hàng đã đăng nhập"],
        ["UC-20", "Xem lịch sử đơn hàng",
         "Danh sách tất cả đơn hàng của khách hàng, sắp xếp theo thời gian giảm dần.",
         "Khách hàng đã đăng nhập"],
        ["UC-21", "Xem chi tiết đơn hàng",
         "Thông tin đầy đủ đơn hàng: danh sách sản phẩm, số lượng, giá, địa chỉ, phương thức thanh toán, trạng thái.",
         "Khách hàng đã đăng nhập"],
        ["UC-22", "Cập nhật trạng thái đơn hàng (Admin)",
         "Admin cập nhật trạng thái: pending → confirmed → shipping → delivered. Có thể hủy đơn: cancelled.",
         "Admin"],
    ],
    caption="Bảng 3.5. Yêu cầu chức năng phân hệ đặt hàng"
)

add_heading("3.2.5. Phân hệ Đánh giá (Review)", level=3)

add_table(
    ["Mã YC", "Tên chức năng", "Mô tả", "Đối tượng"],
    [
        ["UC-23", "Viết đánh giá",
         "Khách hàng đã mua hàng có thể đánh giá sản phẩm với số sao (1-5) và nhận xét văn bản.",
         "Khách hàng đã đăng nhập"],
        ["UC-24", "Xem đánh giá",
         "Hiển thị danh sách đánh giá trên trang chi tiết sản phẩm kèm điểm trung bình.",
         "Tất cả"],
        ["UC-25", "Xoá đánh giá (Admin)",
         "Admin có thể xoá đánh giá không phù hợp.",
         "Admin"],
    ],
    caption="Bảng 3.6. Yêu cầu chức năng phân hệ đánh giá"
)

add_heading("3.2.6. Phân hệ Quản trị (Admin)", level=3)

add_table(
    ["Mã YC", "Tên chức năng", "Mô tả", "Đối tượng"],
    [
        ["UC-26", "Dashboard thống kê",
         "Hiển thị tổng doanh thu, tổng đơn hàng, tổng sản phẩm, tổng người dùng. Biểu đồ doanh thu theo tháng.",
         "Admin"],
        ["UC-27", "Quản lý danh mục",
         "CRUD danh mục sản phẩm (tên, slug, ảnh đại diện).",
         "Admin"],
        ["UC-28", "Quản lý người dùng",
         "Xem danh sách người dùng, kích hoạt/vô hiệu hoá tài khoản.",
         "Admin"],
        ["UC-29", "Upload ảnh",
         "Upload ảnh sản phẩm lên Cloudinary, nhận về URL để gắn vào sản phẩm.",
         "Admin"],
    ],
    caption="Bảng 3.7. Yêu cầu chức năng phân hệ quản trị"
)

add_heading("3.3. Yêu cầu phi chức năng", level=2, numbered="3.3.")

add_paragraph(
    "Ngoài các yêu cầu chức năng, hệ thống cần đáp ứng các yêu cầu phi chức năng "
    "quan trọng sau đây để đảm bảo chất lượng và tính khả dụng trong môi trường thực tế:"
)

add_table(
    ["Loại yêu cầu", "Yêu cầu cụ thể", "Tiêu chí đánh giá"],
    [
        ["Hiệu năng (Performance)",
         "Thời gian phản hồi API trung bình < 200ms cho các query thông thường",
         "Đo bằng công cụ benchmark (Locust/ab)"],
        ["Hiệu năng",
         "Trang sản phẩm load < 3 giây trên kết nối 4G",
         "Google PageSpeed Insights"],
        ["Bảo mật (Security)",
         "Mật khẩu phải hash bằng bcrypt, không lưu plaintext",
         "Kiểm tra database, không thấy password rõ ràng"],
        ["Bảo mật",
         "JWT token có thời hạn, không thể giả mạo",
         "Kiểm tra token giả bị từ chối 401"],
        ["Bảo mật",
         "CORS chỉ cho phép domain được cấu hình",
         "Thử cross-origin request bị chặn"],
        ["Bảo mật",
         "Rate limiting trên endpoint đăng nhập",
         "Quá 10 request/phút → 429 Too Many Requests"],
        ["Khả năng sử dụng (Usability)",
         "Giao diện responsive hoạt động tốt trên mobile, tablet, desktop",
         "Test trên Chrome DevTools Responsive Mode"],
        ["Khả năng sử dụng",
         "Thông báo lỗi rõ ràng và thân thiện với người dùng",
         "Kiểm tra UX manual testing"],
        ["Khả năng bảo trì (Maintainability)",
         "Mã nguồn có cấu trúc rõ ràng, phân lớp (models/schemas/crud/routers)",
         "Code review theo checklist"],
        ["Khả năng mở rộng (Scalability)",
         "Kiến trúc stateless API cho phép scale horizontally",
         "Không lưu state trong server memory"],
        ["Tính sẵn sàng (Availability)",
         "Hệ thống phải tự khởi động lại khi container crash (restart: unless-stopped)",
         "Kiểm tra Docker restart policy"],
        ["Dữ liệu",
         "Dữ liệu PostgreSQL được persist qua Docker volume",
         "Restart container, data vẫn còn nguyên"],
    ],
    caption="Bảng 3.8. Yêu cầu phi chức năng"
)

add_heading("3.4. Biểu đồ Use Case", level=2, numbered="3.4.")

add_heading("3.4.1. Use Case tổng quát", level=3)

add_paragraph(
    "Biểu đồ Use Case tổng quát thể hiện các actor và tập hợp chức năng chính của "
    "hệ thống LuxeBeauty:"
)

add_paragraph(
    "Actor 1 – Khách vãng lai (Guest): Có thể xem danh sách sản phẩm, xem chi tiết "
    "sản phẩm, xem đánh giá, tìm kiếm và lọc sản phẩm. Phải đăng ký hoặc đăng nhập "
    "để thực hiện các chức năng khác.",
    first_indent=False
)

add_paragraph(
    "Actor 2 – Khách hàng đã đăng nhập (Customer): Kế thừa tất cả quyền của Khách vãng "
    "lai, bổ sung thêm: quản lý giỏ hàng, đặt hàng, thanh toán, xem lịch sử đơn hàng, "
    "viết đánh giá, quản lý hồ sơ cá nhân.",
    first_indent=False
)

add_paragraph(
    "Actor 3 – Quản trị viên (Admin): Đăng nhập qua trang admin riêng biệt. Có toàn "
    "quyền quản lý sản phẩm, danh mục, đơn hàng, người dùng và xem báo cáo thống kê.",
    first_indent=False
)

add_heading("3.4.2. Use Case chi tiết – Đặt hàng", level=3)

add_paragraph("Mô tả chi tiết Use Case UC-17 (Đặt hàng):")

add_table(
    ["Thuộc tính", "Nội dung"],
    [
        ["Tên Use Case", "Đặt hàng (Place Order)"],
        ["Mã", "UC-17"],
        ["Actor chính", "Khách hàng đã đăng nhập"],
        ["Điều kiện tiên quyết", "Người dùng đã đăng nhập và có ít nhất 1 sản phẩm trong giỏ hàng"],
        ["Điều kiện hậu", "Đơn hàng mới được tạo trong hệ thống, giỏ hàng được xoá"],
        ["Luồng chính",
         "1. Khách hàng vào trang Checkout\n"
         "2. Hệ thống hiển thị tóm tắt giỏ hàng và form địa chỉ\n"
         "3. Khách hàng điền đầy đủ thông tin giao hàng\n"
         "4. Chọn phương thức thanh toán (PayOS/COD)\n"
         "5. Nhấn 'Đặt hàng'\n"
         "6. Hệ thống validate dữ liệu và kiểm tra tồn kho\n"
         "7. Tạo bản ghi Order và OrderItems trong DB\n"
         "8. Xoá CartItems của user\n"
         "9. Nếu PayOS: tạo payment link và redirect\n"
         "10. Nếu COD: chuyển đến trang thành công"],
        ["Luồng thay thế",
         "6a. Sản phẩm hết hàng: Hiển thị thông báo lỗi, không tạo đơn\n"
         "9a. PayOS lỗi: Hiển thị trang lỗi, đơn hàng ở trạng thái pending"],
        ["Ngoại lệ", "Network timeout: Hiển thị thông báo lỗi kết nối"],
    ],
    caption="Bảng 3.9. Use Case chi tiết – Đặt hàng"
)

add_heading("3.4.3. Use Case chi tiết – Thanh toán PayOS", level=3)

add_table(
    ["Thuộc tính", "Nội dung"],
    [
        ["Tên Use Case", "Thanh toán online qua PayOS"],
        ["Mã", "UC-18"],
        ["Actor chính", "Khách hàng đã đăng nhập, PayOS System"],
        ["Điều kiện tiên quyết", "Đơn hàng đã được tạo thành công với payment_method = 'payos'"],
        ["Luồng chính",
         "1. Backend gọi PayOS API tạo payment link với order_code\n"
         "2. Frontend nhận payment_url và redirect người dùng\n"
         "3. Người dùng quét QR hoặc chuyển khoản trên trang PayOS\n"
         "4. PayOS xác nhận giao dịch thành công\n"
         "5. PayOS gửi POST webhook đến /api/payment/webhook\n"
         "6. Backend verify checksum, cập nhật order status = 'confirmed'\n"
         "7. PayOS redirect người dùng về trang CheckoutSuccess"],
        ["Bảo mật", "Verify HMAC checksum từ PayOS trước khi cập nhật đơn hàng"],
    ],
    caption="Bảng 3.10. Use Case chi tiết – Thanh toán PayOS"
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CHƯƠNG 4 – THIẾT KẾ HỆ THỐNG
# ══════════════════════════════════════════════════════════════════════════════
add_heading("CHƯƠNG 4. THIẾT KẾ HỆ THỐNG", level=1)

add_heading("4.1. Kiến trúc tổng thể", level=2, numbered="4.1.")

add_heading("4.1.1. Kiến trúc phân lớp Backend", level=3)

add_paragraph(
    "Backend của LuxeBeauty được thiết kế theo Clean Architecture với 4 lớp rõ ràng, "
    "đảm bảo tính bảo trì cao và dễ dàng kiểm thử:"
)

add_bullet("Lớp Models (app/models/): Định nghĩa cấu trúc bảng cơ sở dữ liệu bằng "
           "SQLAlchemy ORM. Đây là nguồn dữ liệu duy nhất (Single Source of Truth) về "
           "schema database. Mỗi class Python tương ứng với một bảng SQL.")

add_bullet("Lớp Schemas (app/schemas/): Định nghĩa các Pydantic model cho request/response. "
           "Tách biệt với Models để kiểm soát dữ liệu vào/ra API. Ví dụ: ProductOut "
           "không trả về trường nhạy cảm, UserCreate yêu cầu password tối thiểu 8 ký tự.")

add_bullet("Lớp CRUD (app/crud/): Chứa toàn bộ logic truy vấn database (Create, Read, "
           "Update, Delete). Các hàm CRUD nhận SQLAlchemy Session và trả về Model objects. "
           "Lớp này không biết về HTTP request hay response format.")

add_bullet("Lớp Routers (app/routers/): Xử lý HTTP request/response. Định nghĩa các "
           "endpoint, xác thực authentication/authorization, gọi CRUD functions và "
           "chuyển đổi kết quả sang Pydantic schema.")

add_bullet("Lớp Core (app/core/): Các module cấu hình và tiện ích dùng chung: config.py "
           "(settings từ .env), database.py (engine, session factory), security.py "
           "(JWT, bcrypt), deps.py (FastAPI dependencies).")

add_paragraph(
    "Kiến trúc này đảm bảo nguyên tắc Separation of Concerns: mỗi lớp chỉ biết về "
    "lớp bên dưới nó. Router gọi CRUD, CRUD gọi Models. Thay đổi cách lưu trữ dữ "
    "liệu chỉ cần sửa CRUD mà không ảnh hưởng đến Router, và ngược lại."
)

add_heading("4.1.2. Kiến trúc Frontend", level=3)

add_paragraph(
    "Frontend được tổ chức theo Feature-Based Architecture kết hợp với mô hình "
    "Presentational và Container components:"
)

add_bullet("Pages (src/pages/): Các component trang đầy đủ, tương ứng với các route. "
           "Mỗi page là một container component: fetch data từ API và pass xuống các "
           "presentational components.")

add_bullet("Components (src/components/): Các UI component tái sử dụng (Navbar, Footer, "
           "ProductCard, Layout). Không fetch data trực tiếp, nhận props từ parent.")

add_bullet("Store (src/store/): Zustand stores quản lý global state (auth, cart). "
           "Chỉ lưu state thực sự cần chia sẻ giữa nhiều component.")

add_bullet("API (src/api/): Tập trung toàn bộ logic gọi API. axios.ts cấu hình base URL, "
           "interceptors tự động thêm Authorization header và refresh token khi 401.")

add_bullet("Types (src/types/): TypeScript interfaces dùng chung cho toàn bộ ứng dụng.")

add_heading("4.2. Thiết kế Cơ sở dữ liệu", level=2, numbered="4.2.")

add_heading("4.2.1. Mô hình thực thể quan hệ (ERD)", level=3)

add_paragraph(
    "Cơ sở dữ liệu PostgreSQL của LuxeBeauty gồm 7 bảng chính với mối quan hệ như sau: "
    "users là trung tâm, liên kết với cart_items, orders và reviews. products liên kết "
    "với categories, cart_items, order_items và reviews. orders liên kết với order_items."
)

add_heading("4.2.2. Mô tả chi tiết các bảng", level=3)

add_paragraph("Bảng users – Lưu thông tin tài khoản người dùng:")

add_table(
    ["Tên cột", "Kiểu dữ liệu", "Ràng buộc", "Mô tả"],
    [
        ["id", "UUID", "PK, DEFAULT gen_random_uuid()", "Khoá chính tự sinh"],
        ["email", "VARCHAR(255)", "NOT NULL, UNIQUE", "Email đăng nhập"],
        ["password_hash", "VARCHAR(255)", "NOT NULL", "Mật khẩu đã hash (bcrypt)"],
        ["full_name", "VARCHAR(255)", "NOT NULL", "Họ và tên đầy đủ"],
        ["phone", "VARCHAR(20)", "NULLABLE", "Số điện thoại"],
        ["role", "VARCHAR(20)", "DEFAULT 'customer'", "Vai trò: customer/admin"],
        ["is_active", "BOOLEAN", "DEFAULT true", "Tài khoản kích hoạt hay bị khóa"],
        ["created_at", "TIMESTAMP", "DEFAULT NOW()", "Thời gian tạo tài khoản"],
    ],
    caption="Bảng 4.1. Cấu trúc bảng users"
)

add_paragraph("Bảng products – Lưu thông tin sản phẩm:")

add_table(
    ["Tên cột", "Kiểu dữ liệu", "Ràng buộc", "Mô tả"],
    [
        ["id", "UUID", "PK", "Khoá chính"],
        ["name", "VARCHAR(500)", "NOT NULL", "Tên sản phẩm"],
        ["slug", "VARCHAR(500)", "NOT NULL, UNIQUE", "URL-friendly name (auto-gen)"],
        ["description", "TEXT", "NULLABLE", "Mô tả chi tiết sản phẩm"],
        ["price", "DECIMAL(12,2)", "NOT NULL", "Giá gốc"],
        ["sale_price", "DECIMAL(12,2)", "NULLABLE", "Giá khuyến mãi (có thể null)"],
        ["stock", "INTEGER", "NOT NULL, DEFAULT 0", "Số lượng tồn kho"],
        ["images", "ARRAY(TEXT)", "DEFAULT '{}'", "Mảng URL ảnh sản phẩm"],
        ["category_id", "UUID", "FK → categories.id", "Danh mục sản phẩm"],
        ["brand", "VARCHAR(255)", "NULLABLE", "Tên thương hiệu"],
        ["is_active", "BOOLEAN", "DEFAULT true", "Sản phẩm đang kinh doanh"],
        ["created_at", "TIMESTAMP", "DEFAULT NOW()", "Thời gian tạo"],
    ],
    caption="Bảng 4.2. Cấu trúc bảng products"
)

add_paragraph("Bảng orders – Lưu thông tin đơn hàng:")

add_table(
    ["Tên cột", "Kiểu dữ liệu", "Ràng buộc", "Mô tả"],
    [
        ["id", "UUID", "PK", "Khoá chính"],
        ["user_id", "UUID", "FK → users.id", "Người đặt hàng"],
        ["status", "VARCHAR(20)", "DEFAULT 'pending'", "Trạng thái đơn hàng"],
        ["total_price", "DECIMAL(12,2)", "NOT NULL", "Tổng tiền đơn hàng"],
        ["shipping_address", "JSONB", "NOT NULL", "Địa chỉ giao hàng (JSON)"],
        ["payment_method", "VARCHAR(20)", "NOT NULL", "Phương thức: payos/cod"],
        ["order_code", "VARCHAR(50)", "UNIQUE, NULLABLE", "Mã đơn hàng (cho PayOS)"],
        ["created_at", "TIMESTAMP", "DEFAULT NOW()", "Thời gian đặt hàng"],
    ],
    caption="Bảng 4.3. Cấu trúc bảng orders"
)

add_paragraph(
    "Trạng thái đơn hàng (status) có thể nhận các giá trị: pending (chờ xác nhận), "
    "confirmed (đã xác nhận), shipping (đang giao), delivered (đã giao), cancelled (đã hủy). "
    "Luồng chuyển trạng thái thông thường: pending → confirmed → shipping → delivered."
)

add_paragraph("Bảng order_items – Chi tiết từng sản phẩm trong đơn hàng:")

add_table(
    ["Tên cột", "Kiểu dữ liệu", "Ràng buộc", "Mô tả"],
    [
        ["id", "UUID", "PK", "Khoá chính"],
        ["order_id", "UUID", "FK → orders.id", "Thuộc đơn hàng nào"],
        ["product_id", "UUID", "FK → products.id", "Sản phẩm nào"],
        ["quantity", "INTEGER", "NOT NULL", "Số lượng đặt"],
        ["price_at_purchase", "DECIMAL(12,2)", "NOT NULL",
         "Giá tại thời điểm mua (snapshot)"],
    ],
    caption="Bảng 4.4. Cấu trúc bảng order_items"
)

add_paragraph(
    "Trường price_at_purchase là quan trọng: nó lưu giá sản phẩm tại thời điểm mua, "
    "không tham chiếu đến giá hiện tại. Điều này đảm bảo lịch sử đơn hàng không bị "
    "ảnh hưởng khi admin thay đổi giá sản phẩm."
)

add_paragraph("Bảng cart_items – Giỏ hàng tạm thời:")

add_table(
    ["Tên cột", "Kiểu dữ liệu", "Ràng buộc", "Mô tả"],
    [
        ["id", "UUID", "PK", "Khoá chính"],
        ["user_id", "UUID", "FK → users.id", "Chủ giỏ hàng"],
        ["product_id", "UUID", "FK → products.id", "Sản phẩm trong giỏ"],
        ["quantity", "INTEGER", "NOT NULL, DEFAULT 1", "Số lượng"],
    ],
    caption="Bảng 4.5. Cấu trúc bảng cart_items"
)

add_paragraph("Bảng reviews – Đánh giá sản phẩm:")

add_table(
    ["Tên cột", "Kiểu dữ liệu", "Ràng buộc", "Mô tả"],
    [
        ["id", "UUID", "PK", "Khoá chính"],
        ["user_id", "UUID", "FK → users.id", "Người đánh giá"],
        ["product_id", "UUID", "FK → products.id", "Sản phẩm được đánh giá"],
        ["rating", "INTEGER", "CHECK(1-5)", "Số sao (1 đến 5)"],
        ["comment", "TEXT", "NULLABLE", "Nhận xét bằng văn bản"],
        ["created_at", "TIMESTAMP", "DEFAULT NOW()", "Thời gian đánh giá"],
    ],
    caption="Bảng 4.6. Cấu trúc bảng reviews"
)

add_paragraph("Bảng categories – Danh mục sản phẩm:")

add_table(
    ["Tên cột", "Kiểu dữ liệu", "Ràng buộc", "Mô tả"],
    [
        ["id", "UUID", "PK", "Khoá chính"],
        ["name", "VARCHAR(255)", "NOT NULL", "Tên danh mục"],
        ["slug", "VARCHAR(255)", "NOT NULL, UNIQUE", "URL-friendly name"],
        ["image_url", "TEXT", "NULLABLE", "Ảnh đại diện danh mục"],
    ],
    caption="Bảng 4.7. Cấu trúc bảng categories"
)

add_heading("4.3. Thiết kế API", level=2, numbered="4.3.")

add_heading("4.3.1. Chuẩn API RESTful", level=3)

add_paragraph(
    "Toàn bộ API của LuxeBeauty tuân thủ thiết kế RESTful với các quy tắc sau:"
)

add_bullet("URL Resources: Danh từ số nhiều (/api/products, /api/orders, /api/users).")
add_bullet("HTTP Methods: GET (đọc), POST (tạo mới), PUT (cập nhật toàn bộ), "
           "PATCH (cập nhật một phần), DELETE (xoá).")
add_bullet("Status Codes: 200 OK, 201 Created, 400 Bad Request, 401 Unauthorized, "
           "403 Forbidden, 404 Not Found, 422 Unprocessable Entity, 429 Too Many Requests, "
           "500 Internal Server Error.")
add_bullet("Base URL: Tất cả API endpoint có prefix /api/.")
add_bullet("Content Type: application/json cho request body và response.")
add_bullet("Versioning: Hiện tại chưa có versioning (v1), có thể bổ sung trong tương lai.")

add_heading("4.3.2. Danh sách API Endpoints", level=3)

add_table(
    ["Method", "Endpoint", "Mô tả", "Auth"],
    [
        ["POST", "/api/auth/register", "Đăng ký tài khoản mới", "Không"],
        ["POST", "/api/auth/login", "Đăng nhập, nhận JWT token", "Không"],
        ["POST", "/api/auth/refresh", "Làm mới access token", "Refresh token"],
        ["GET", "/api/products", "Danh sách sản phẩm (filter, search, sort, page)", "Không"],
        ["GET", "/api/products/{slug}", "Chi tiết sản phẩm theo slug", "Không"],
        ["POST", "/api/products", "Tạo sản phẩm mới", "Admin"],
        ["PUT", "/api/products/{id}", "Cập nhật sản phẩm", "Admin"],
        ["DELETE", "/api/products/{id}", "Xoá sản phẩm", "Admin"],
        ["GET", "/api/categories", "Danh sách danh mục", "Không"],
        ["POST", "/api/categories", "Tạo danh mục mới", "Admin"],
        ["PUT", "/api/categories/{id}", "Cập nhật danh mục", "Admin"],
        ["DELETE", "/api/categories/{id}", "Xoá danh mục", "Admin"],
        ["GET", "/api/cart", "Lấy giỏ hàng của user hiện tại", "Customer"],
        ["POST", "/api/cart", "Thêm sản phẩm vào giỏ", "Customer"],
        ["PUT", "/api/cart/{item_id}", "Cập nhật số lượng trong giỏ", "Customer"],
        ["DELETE", "/api/cart/{item_id}", "Xoá sản phẩm khỏi giỏ", "Customer"],
        ["POST", "/api/orders", "Tạo đơn hàng từ giỏ hàng", "Customer"],
        ["GET", "/api/orders", "Lịch sử đơn hàng của user", "Customer"],
        ["GET", "/api/orders/{id}", "Chi tiết đơn hàng", "Customer"],
        ["PUT", "/api/orders/{id}/status", "Cập nhật trạng thái đơn hàng", "Admin"],
        ["GET", "/api/users/me", "Thông tin profile user hiện tại", "Customer"],
        ["PUT", "/api/users/me", "Cập nhật profile", "Customer"],
        ["GET", "/api/admin/stats", "Thống kê tổng quan dashboard", "Admin"],
        ["GET", "/api/admin/products", "Danh sách sản phẩm (admin view)", "Admin"],
        ["GET", "/api/admin/orders", "Danh sách tất cả đơn hàng", "Admin"],
        ["GET", "/api/admin/users", "Danh sách tất cả người dùng", "Admin"],
        ["POST", "/api/upload/image", "Upload ảnh lên Cloudinary", "Admin"],
        ["POST", "/api/payment/webhook", "Nhận webhook từ PayOS", "PayOS HMAC"],
        ["POST", "/api/payment/simulate-success", "Giả lập thanh toán thành công", "Customer"],
        ["GET", "/api/health", "Health check endpoint", "Không"],
    ],
    caption="Bảng 4.8. Danh sách đầy đủ API Endpoints"
)

add_heading("4.3.3. Ví dụ Request/Response", level=3)

add_paragraph("Ví dụ 1: Đăng ký tài khoản mới")

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.left_indent = Cm(1.0)
run = p.add_run("POST /api/auth/register\n"
                "Content-Type: application/json\n\n"
                "Request Body:\n"
                '{\n'
                '  "email": "nguyenvan@gmail.com",\n'
                '  "password": "MyPass@2026",\n'
                '  "full_name": "Nguyễn Văn Sơn"\n'
                '}\n\n'
                "Response 201 Created:\n"
                '{\n'
                '  "id": "550e8400-e29b-41d4-a716-446655440000",\n'
                '  "email": "nguyenvan@gmail.com",\n'
                '  "full_name": "Nguyễn Văn Sơn",\n'
                '  "role": "customer",\n'
                '  "is_active": true\n'
                '}')
set_font(run, name="Courier New", size=10)

add_paragraph("Ví dụ 2: Lấy danh sách sản phẩm với filter")

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.left_indent = Cm(1.0)
run = p.add_run("GET /api/products?category=skincare&min_price=100000&max_price=500000\n"
                "         &search=serum&sort=price_asc&page=1&page_size=20\n\n"
                "Response 200 OK:\n"
                '{\n'
                '  "items": [\n'
                '    {\n'
                '      "id": "...",\n'
                '      "name": "Serum Dưỡng Da COSRX",\n'
                '      "slug": "serum-duong-da-cosrx",\n'
                '      "price": 350000,\n'
                '      "sale_price": 299000,\n'
                '      "images": ["https://res.cloudinary.com/..."],\n'
                '      "brand": "COSRX",\n'
                '      "stock": 50\n'
                '    }\n'
                '  ],\n'
                '  "total": 15,\n'
                '  "page": 1,\n'
                '  "page_size": 20\n'
                '}')
set_font(run, name="Courier New", size=10)

add_heading("4.4. Thiết kế Giao diện", level=2, numbered="4.4.")

add_heading("4.4.1. Nguyên tắc thiết kế UI", level=3)

add_paragraph(
    "Giao diện LuxeBeauty được thiết kế theo phong cách hiện đại, sang trọng phù hợp "
    "với thương hiệu mỹ phẩm cao cấp. Các nguyên tắc thiết kế được áp dụng:"
)

add_bullet("Bảng màu: Màu nền trắng (#FFFFFF) và xám nhạt (#F9FAFB) tạo cảm giác sạch sẽ, "
           "tinh tế. Màu chủ đạo: hồng đậm (#DB2777, pink-600) và đen (#111827) tạo điểm "
           "nhấn sang trọng.")
add_bullet("Typography: Font chữ Inter (sans-serif) cho giao diện, kết hợp với cỡ chữ "
           "phân cấp rõ ràng (heading 2xl-4xl, body text base-lg).")
add_bullet("Responsive Design: Sử dụng Tailwind CSS grid system. Mobile (1 cột) → "
           "Tablet (2 cột) → Desktop (3-4 cột) cho danh sách sản phẩm.")
add_bullet("Whitespace: Sử dụng padding và margin rộng rãi (p-6, p-8) tạo cảm giác "
           "thoải mái, không chật chội.")
add_bullet("Micro-interactions: Framer Motion animation cho card hover, page transition, "
           "toast notification (react-hot-toast).")
add_bullet("Accessibility: Alt text cho ảnh, label cho form inputs, contrast ratio đủ "
           "theo WCAG 2.1 AA.")

add_heading("4.4.2. Cấu trúc trang chính", level=3)

add_table(
    ["Trang", "URL", "Thành phần chính", "Mô tả"],
    [
        ["Trang chủ", "/", "Hero banner, sản phẩm nổi bật, danh mục, CTA",
         "Landing page với banner quảng cáo và sản phẩm được chọn lọc"],
        ["Danh sách sản phẩm", "/products", "Sidebar filter, product grid, pagination",
         "Danh sách sản phẩm có thể lọc theo nhiều tiêu chí"],
        ["Chi tiết sản phẩm", "/products/:slug", "Image gallery, thông tin, nút thêm giỏ, reviews",
         "Trang chi tiết với ảnh carousel và đánh giá của khách hàng"],
        ["Giỏ hàng", "/cart", "Danh sách cart items, tổng tiền, nút checkout",
         "Quản lý sản phẩm trong giỏ trước khi đặt hàng"],
        ["Thanh toán", "/checkout", "Form địa chỉ, tóm tắt đơn, chọn thanh toán",
         "Form nhập địa chỉ và xác nhận đơn hàng"],
        ["Đăng nhập/Đăng ký", "/login, /register", "Form xác thực",
         "Trang xác thực người dùng với validation realtime"],
        ["Tài khoản", "/account", "Thông tin cá nhân, lịch sử đơn hàng",
         "Quản lý hồ sơ và xem lịch sử mua sắm"],
        ["Admin Dashboard", "/admin", "Sidebar navigation, stats cards, charts",
         "Trang tổng quan với biểu đồ doanh thu (Recharts)"],
        ["Admin Sản phẩm", "/admin/products", "Table CRUD, modal form",
         "Quản lý toàn bộ danh mục sản phẩm"],
        ["Admin Đơn hàng", "/admin/orders", "Table với filter trạng thái",
         "Xem và cập nhật trạng thái đơn hàng"],
    ],
    caption="Bảng 4.9. Cấu trúc các trang giao diện"
)

add_heading("4.5. Luồng xử lý chính", level=2, numbered="4.5.")

add_heading("4.5.1. Luồng xác thực người dùng", level=3)

add_paragraph("Luồng đăng nhập và quản lý JWT token:")

add_paragraph(
    "Bước 1: Người dùng nhập email và mật khẩu trên trang đăng nhập.\n"
    "Bước 2: Frontend gửi POST /api/auth/login với {email, password}.\n"
    "Bước 3: Backend lấy user từ DB theo email, verify password hash bằng bcrypt.\n"
    "Bước 4: Nếu hợp lệ, tạo Access Token (30 phút) và Refresh Token (7 ngày) bằng python-jose.\n"
    "Bước 5: Trả về {access_token, refresh_token, token_type: 'bearer'}.\n"
    "Bước 6: Frontend lưu tokens vào Zustand store (và localStorage để persist qua refresh trang).\n"
    "Bước 7: Axios interceptor tự động thêm header: Authorization: Bearer {access_token} cho mọi request tiếp theo.\n"
    "Bước 8: Khi nhận 401 Unauthorized, interceptor tự động gọi /auth/refresh để lấy token mới.",
    first_indent=False
)

add_heading("4.5.2. Luồng mua hàng đầy đủ", level=3)

add_paragraph("Luồng từ duyệt sản phẩm đến hoàn tất đặt hàng:")

add_paragraph(
    "Bước 1 – Duyệt sản phẩm: Khách hàng truy cập /products, chọn lọc theo danh mục "
    "'Chăm sóc da', xem danh sách. React Query gọi GET /api/products với query params "
    "và cache kết quả 5 phút.",
    first_indent=False
)

add_paragraph(
    "Bước 2 – Xem chi tiết: Nhấp vào sản phẩm → chuyển đến /products/serum-duong-da-cosrx. "
    "Trang fetch GET /api/products/serum-duong-da-cosrx. Hiển thị gallery ảnh, mô tả, "
    "giá, và reviews.",
    first_indent=False
)

add_paragraph(
    "Bước 3 – Thêm vào giỏ: Nhấn nút 'Thêm vào giỏ'. Nếu chưa đăng nhập → redirect "
    "đến /login với returnUrl. Nếu đã đăng nhập → gọi POST /api/cart, backend kiểm tra "
    "tồn kho, thêm cart_item vào DB, cập nhật cartStore. Icon giỏ hàng trên Navbar "
    "hiện số lượng mới.",
    first_indent=False
)

add_paragraph(
    "Bước 4 – Thanh toán: Vào /cart, xem lại giỏ hàng, nhấn 'Tiến hành thanh toán'. "
    "Chuyển đến /checkout, điền địa chỉ giao hàng. Chọn 'Thanh toán qua PayOS'. "
    "Nhấn 'Đặt hàng'.",
    first_indent=False
)

add_paragraph(
    "Bước 5 – Xử lý đặt hàng: Frontend gọi POST /api/orders với {shipping_address, "
    "payment_method: 'payos'}. Backend: (a) Lấy cart items của user, (b) Tính total_price, "
    "(c) Tạo Order record trong DB, (d) Tạo OrderItem records với price_at_purchase, "
    "(e) Xoá CartItems, (f) Gọi PayOS API tạo payment link, (g) Trả về {order_id, payment_url}.",
    first_indent=False
)

add_paragraph(
    "Bước 6 – Thanh toán PayOS: Frontend redirect đến payment_url của PayOS. Khách hàng "
    "quét QR Code trên app ngân hàng. PayOS xử lý giao dịch và gửi POST webhook đến "
    "/api/payment/webhook với HMAC signature. Backend verify signature, cập nhật "
    "order.status = 'confirmed'. PayOS redirect về /checkout/success?orderCode=xxx.",
    first_indent=False
)

add_paragraph(
    "Bước 7 – Hoàn tất: Trang CheckoutSuccess hiển thị thông tin đơn hàng thành công. "
    "Khách hàng có thể xem lại đơn hàng tại /account.",
    first_indent=False
)

add_heading("4.5.3. Luồng quản lý sản phẩm (Admin)", level=3)

add_paragraph(
    "Admin muốn thêm sản phẩm mới thực hiện qua các bước:"
)

add_bullet("Đăng nhập admin tại /admin/login với tài khoản admin@luxebeauty.vn.")
add_bullet("Truy cập /admin/products → Nhấn 'Thêm sản phẩm mới'.")
add_bullet("Nhập thông tin: tên sản phẩm (slug tự động tạo), mô tả, giá, giá sale, tồn kho, thương hiệu, danh mục.")
add_bullet("Upload ảnh: Nhấn nút Upload, chọn file ảnh. Frontend gọi POST /api/upload/image. "
           "Backend nhận file, upload lên Cloudinary, trả về URL. URL được thêm vào danh sách ảnh sản phẩm.")
add_bullet("Nhấn 'Lưu' → Frontend gọi POST /api/products với đầy đủ thông tin. "
           "Backend validate, tạo product record trong DB, trả về ProductOut.")
add_bullet("Sản phẩm xuất hiện ngay trong danh sách và có thể mua trên trang khách hàng.")

doc.add_page_break()

doc.save(OUT)
print(f"Saved: Chương 3 + 4 xong → {OUT}")
